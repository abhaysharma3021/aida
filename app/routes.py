from models.groq_client import GroqClient
import os
import sys
import json
from venv import logger
from flask import Blueprint, render_template, request, flash, redirect, url_for, session, send_file, current_app
from datetime import datetime
import markdown
import io
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from difflib import get_close_matches
import zipfile
from models.course_materials import CourseMaterialsGenerator, generate_course_materials
from collections import defaultdict
import html
from xml.etree import ElementTree as ET
from xml.dom import minidom
from functools import wraps
from app.analysis_log import AnalysisLog
import time





# Create the Blueprint
main = Blueprint('main', __name__)

# Define the form classes
from flask_wtf import FlaskForm
from wtforms import StringField, TextAreaField, SelectField, SubmitField, IntegerField
from wtforms.validators import DataRequired, Optional, Length

class CourseInputForm(FlaskForm):
    course_topic = StringField('Course Topic', validators=[DataRequired()])
    audience_type = SelectField('Audience Type', 
                              choices=[('beginner', 'Beginner'), 
                                      ('intermediate', 'Intermediate'),
                                      ('advanced', 'Advanced')],
                              validators=[DataRequired()])
    job_titles = TextAreaField('Role of Targeted Audience', 
                             description="Enter one job title per line (e.g., Software Engineer, Project Manager)")
    submit = SubmitField('Continue to Audience Analysis')

class TerminalObjectivesForm(FlaskForm):
    terminal_objectives = TextAreaField('Terminal Objectives', 
                                       description="Enter the terminal objectives for this course",
                                       validators=[DataRequired()])
    submit = SubmitField('Generate Task Analysis')

class CourseDesignRequirementsForm(FlaskForm):
    """Form for collecting additional requirements for course design."""
    course_duration = StringField('Course Duration (e.g., 3 days, 8 weeks)', validators=[Optional()])
    delivery_format = SelectField('Delivery Format', 
                                 choices=[
                                     ('', 'Select Format'),
                                     ('in_person', 'In-Person'),
                                     ('online_synchronous', 'Online Synchronous'),
                                     ('online_asynchronous', 'Online Asynchronous'),
                                     ('blended', 'Blended/Hybrid')
                                 ],
                                 validators=[Optional()])
    module_count = IntegerField('Number of Modules (3-7 recommended)', validators=[Optional()])
    additional_requirements = TextAreaField('Additional Requirements or Notes', 
                                          validators=[Optional(), Length(max=2000)])
    submit = SubmitField('Generate Course Design')

# Import models with absolute path
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from models.audience_analysis import generate_audience_analysis
from models.task_analysis import generate_task_analysis
from models.course_design import generate_course_structure, generate_instructional_strategies, generate_assessment_plan, generate_comprehensive_course_design

# Create a data directory if it doesn't exist
DATA_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'data')
if not os.path.exists(DATA_DIR):
    os.makedirs(DATA_DIR)

def login_required(view):
    @wraps(view)
    def wrapped(*args, **kwargs):
        if "user" not in session:
            return redirect(url_for("auth.login"))
        return view(*args, **kwargs)
    return wrapped

def save_analysis(analysis_id, data):
    """Save analysis data to a JSON file"""
    filepath = os.path.join(DATA_DIR, f"{analysis_id}.json")
    with open(filepath, 'w') as f:
        json.dump(data, f)
    return analysis_id

def load_analysis(analysis_id):
    """Load analysis data from a JSON file"""
    filepath = os.path.join(DATA_DIR, f"{analysis_id}.json")
    if os.path.exists(filepath):
        with open(filepath, 'r') as f:
            return json.load(f)
    return None

def markdown_to_docx(doc, markdown_text):
    """Convert markdown text to formatted docx document."""
    # Split the markdown into lines
    lines = markdown_text.split('\n')
    current_list = None
    list_item_pattern = re.compile(r'^\s*[-*]\s+(.+)$')
    heading_pattern = re.compile(r'^(#{1,6})\s+(.+)$')
    task_heading_pattern = re.compile(r'^\*\*(Task\s+\d+:.+?)\*\*$')
    learning_activity_pattern = re.compile(r'^\*\*(Learning Activity:)\*\*\s*(.+)$')
    assessment_pattern = re.compile(r'^\*\*(Assessment:)\*\*\s*(.+)$')
    
    for line in lines:
        if not line.strip():
            # Empty line
            doc.add_paragraph()
            current_list = None
            continue
        
        # Check if line is a heading
        heading_match = heading_pattern.match(line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            doc.add_heading(text, level)
            current_list = None
            continue
            
        # Check for task headings like "**Task 1: Understanding Basic Syntax**"
        task_match = task_heading_pattern.match(line)
        if task_match:
            title = task_match.group(1)
            heading = doc.add_heading(title, 2)
            current_list = None
            continue
            
        # Check for "**Learning Activity:**" or "**Assessment:**"
        learning_match = learning_activity_pattern.match(line)
        if learning_match:
            p = doc.add_paragraph()
            run = p.add_run(learning_match.group(1) + " ")
            run.bold = True
            if learning_match.group(2):
                p.add_run(learning_match.group(2))
            current_list = None
            continue
            
        assessment_match = assessment_pattern.match(line)
        if assessment_match:
            p = doc.add_paragraph()
            run = p.add_run(assessment_match.group(1) + " ")
            run.bold = True
            if assessment_match.group(2):
                p.add_run(assessment_match.group(2))
            current_list = None
            continue
        
        # Check if line is a list item
        list_match = list_item_pattern.match(line)
        if list_match:
            if current_list is None:
                current_list = doc.add_paragraph(style='List Bullet')
            else:
                current_list = doc.add_paragraph(style='List Bullet')
                
            # Process the list item text for any formatting
            item_text = list_match.group(1)
            remaining_text = item_text
            
            # Process bold and italic in list item
            while '**' in remaining_text or '*' in remaining_text:
                # Check for bold text
                bold_match = re.search(r'\*\*(.+?)\*\*', remaining_text)
                if bold_match:
                    start, end = bold_match.span()
                    # Add text before the bold part
                    if start > 0:
                        current_list.add_run(remaining_text[:start])
                    # Add the bold text
                    current_list.add_run(bold_match.group(1)).bold = True
                    # Update remaining text
                    remaining_text = remaining_text[end:]
                    continue
                
                # Check for italic text
                italic_match = re.search(r'\*(.+?)\*', remaining_text)
                if italic_match:
                    start, end = italic_match.span()
                    # Add text before the italic part
                    if start > 0:
                        current_list.add_run(remaining_text[:start])
                    # Add the italic text
                    current_list.add_run(italic_match.group(1)).italic = True
                    # Update remaining text
                    remaining_text = remaining_text[end:]
                    continue
                
                break
            
            # Add any remaining text
            if remaining_text:
                current_list.add_run(remaining_text)
            
            continue
        
        # Regular paragraph - check for special patterns
        if '**' in line:
            p = doc.add_paragraph()
            remaining_text = line
            
            # Process bold and italic formatting
            while '**' in remaining_text or '*' in remaining_text:
                # Check for bold text
                bold_match = re.search(r'\*\*(.+?)\*\*', remaining_text)
                if bold_match:
                    start, end = bold_match.span()
                    # Add text before the bold part
                    if start > 0:
                        p.add_run(remaining_text[:start])
                    # Add the bold text
                    p.add_run(bold_match.group(1)).bold = True
                    # Update remaining text
                    remaining_text = remaining_text[end:]
                    continue
                
                # Check for italic text
                italic_match = re.search(r'\*(.+?)\*', remaining_text)
                if italic_match:
                    start, end = italic_match.span()
                    # Add text before the italic part
                    if start > 0:
                        p.add_run(remaining_text[:start])
                    # Add the italic text
                    p.add_run(italic_match.group(1)).italic = True
                    # Update remaining text
                    remaining_text = remaining_text[end:]
                    continue
                
                break
            
            # Add any remaining text
            if remaining_text:
                p.add_run(remaining_text)
        else:
            # Regular paragraph with no special formatting
            p = doc.add_paragraph(line)
        
        current_list = None
def escape_control_characters(json_str):
    def replace_control_chars(match):
        # Get the matched string (content between quotes)
        content = match.group(0)
        # Replace \r\n with \\r\\n and \n with \\n
        content = content.replace('\r\n', '\\r\\n').replace('\n', '\\n')
        return content
    
    # Match string literals (text between double quotes)
    pattern = r'"[^"]*"'
    return re.sub(pattern, replace_control_chars, json_str)

# Function to unescape control characters in the extracted content
def unescape_content(content):
    # Convert literal \r\n and \n to actual control characters
    return content.encode().decode('unicode_escape')

@main.route('/', methods=['GET', 'POST'])
@login_required
def index():
    form = CourseInputForm()
    if form.validate_on_submit():
        try:
            # Process form data
            course_topic = form.course_topic.data
            audience_type = form.audience_type.data
            job_titles = form.job_titles.data
            
            current_app.logger.info("Received form data: course_topic=%s, audience_type=%s", course_topic, audience_type)

            # Create a timestamp-based ID
            timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
            analysis_id = f"analysis_{timestamp}"
            
            # Generate audience analysis
            audience_analysis = generate_audience_analysis(course_topic, audience_type, '')
            current_app.logger.debug("Generated audience analysis for ID: %s", analysis_id)

            
            # Save initial data to file
            analysis_data = {
                'audience_analysis': audience_analysis,
                'course_topic': course_topic,
                'audience_type': audience_type,
                'job_titles': job_titles,
                'generated_date': datetime.now().strftime("%B %d, %Y at %H:%M")
            }
            save_analysis(analysis_id, analysis_data)
            current_app.logger.info("Analysis data saved with ID: %s", analysis_id)
            
            logged_user = session.get("user")
            if not logged_user:
                flash("User session not found.")
                return redirect(url_for("auth_bp.login"))  # or any fallback
             
            email = logged_user["email"]
            
            AnalysisLog.create(
                useremail= email,
                analysis_id=analysis_id,
                data= analysis_data
            )
            
            # Store the ID in session
            session['current_analysis_id'] = analysis_id
            current_app.logger.debug("Stored analysis ID in session: %s", analysis_id)
            
            return redirect(url_for('main.audience_analysis', analysis_id=analysis_id))
        
        except Exception as e:
            current_app.logger.error("Error generating analysis: %s", str(e), exc_info=True)
            flash(f"Error generating analysis: {str(e)}")
            print(f"Error details: {str(e)}")
            return render_template('index.html', form=form)
        
    current_app.logger.debug("Rendering form without submission.")
    return render_template('index.html', form=form)

@main.route('/audience_analysis/<analysis_id>', methods=['GET', 'POST'])
def audience_analysis(analysis_id):

    current_app.logger.info("Accessed audience_analysis view for ID: %s", analysis_id)
    # Load analysis data from file
    analysis_data = load_analysis(analysis_id)
    
    if not analysis_data:
        current_app.logger.warning("Analysis ID %s not found. Redirecting to index.", analysis_id)
        flash('Analysis not found. Please submit the form to generate a new analysis.')
        return redirect(url_for('main.index'))
    
    current_app.logger.debug("Loaded analysis data for ID: %s", analysis_id)

    if request.method == 'POST':
        # Check if this is a task analysis generation request
        if request.form.get('generate_task_analysis') == 'true':
            try:
                print(f"Generating task analysis for {analysis_data['course_topic']}")  # Debug log
                
                # Generate task analysis without terminal objectives
                task_analysis = generate_task_analysis(
                    analysis_data['course_topic'], 
                    analysis_data['audience_type'],
                    ""  # No terminal objectives needed
                )
                
                print("Task analysis generated successfully")  # Debug log
                
                # Add task analysis to the data
                analysis_data['task_analysis'] = task_analysis
                
                # Save updated analysis
                save_analysis(analysis_id, analysis_data)
                
                print(f"Redirecting to task_analysis page for {analysis_id}")  # Debug log
                
                flash('Task analysis generated successfully!')
                return redirect(url_for('main.task_analysis', analysis_id=analysis_id))
                
            except Exception as e:
                print(f"Error generating task analysis: {str(e)}")  # Debug log
                flash(f'Error generating task analysis: {str(e)}')
                # Fall through to render the same page with error
        
        # If we get here, it's likely the old terminal objectives form submission
        # You can remove this section if you don't need it anymore
        else:
            # Handle any other POST requests (like old terminal objectives form)
            flash('Invalid form submission.')
                                                                
    # Initialize the terminal objectives form
    form = TerminalObjectivesForm()
    
    # Pre-populate the form if terminal objectives exist
    if 'terminal_objectives' in analysis_data:
        form.terminal_objectives.data = analysis_data['terminal_objectives']
        current_app.logger.debug("Pre-populated terminal objectives for ID: %s", analysis_id)

    
    if form.validate_on_submit():

        try:
            # Get terminal objectives from form
            terminal_objectives = form.terminal_objectives.data
            
            # Update the analysis data
            analysis_data['terminal_objectives'] = terminal_objectives

            current_app.logger.info("Received terminal objectives for ID: %s", analysis_id)

            # Generate task analysis
            task_analysis = generate_task_analysis(
                analysis_data['course_topic'], 
                analysis_data['audience_type'],
                terminal_objectives
            )
            
            # Add task analysis to the data
            analysis_data['task_analysis'] = task_analysis
            
            # Save updated analysis
            save_analysis(analysis_id, analysis_data)
            current_app.logger.info("Saved updated analysis with task analysis for ID: %s", analysis_id)

            
            # AnalysisLog.update_by_analysis_id(
            #     analysis_id=analysis_id,
            #     data=analysis_data
            # )
            
            # Redirect to task analysis page
            return redirect(url_for('main.task_analysis', analysis_id=analysis_id))
        
        except Exception as e:
            current_app.logger.error("Error processing terminal objectives for ID %s: %s", analysis_id, str(e), exc_info=True)
            flash(f"An error occurred: {str(e)}")
            return render_template('audience_analysis.html',
                                   audience_analysis='',
                                   course_topic='',
                                   current_date='',
                                   analysis_id=analysis_id,
                                   form=form,
                                   analysis_data={})

    # Convert Markdown to HTML
    audience_analysis_html = markdown.markdown(analysis_data['audience_analysis'])
    current_app.logger.debug("Rendered Markdown for audience analysis ID: %s", analysis_id)

    # For GET request, display the audience analysis with terminal objectives form
    return render_template('audience_analysis.html', 
                          audience_analysis=audience_analysis_html,
                          course_topic=analysis_data['course_topic'],
                          current_date=analysis_data['generated_date'],
                          analysis_id=analysis_id,
                          form=form,
                          analysis_data=analysis_data)  # Pass the full analysis_data

@main.route('/task_analysis/<analysis_id>')
def task_analysis(analysis_id):
    # Load analysis data from file
    analysis_data = load_analysis(analysis_id)
    
    if not analysis_data:
        flash('Analysis not found. Please submit the form to generate a new analysis.')
        return redirect(url_for('main.index'))
    
    # Convert Markdown to HTML
    try:
        # Convert Markdown to HTML
        task_analysis_html = markdown.markdown(analysis_data['task_analysis'])
        current_app.logger.debug("Converted task analysis markdown to HTML for ID: %s", analysis_id)
    except Exception as e:
        current_app.logger.error("Error converting task analysis markdown for ID %s: %s", analysis_id, str(e), exc_info=True)
        flash("Failed to render task analysis.")
        return redirect(url_for('main.index'))
    
    return render_template('task_analysis.html', 
                          task_analysis=task_analysis_html,
                          course_topic=analysis_data['course_topic'],
                          current_date=analysis_data['generated_date'],
                          analysis_id=analysis_id)

@main.route('/prepare_course_design/<analysis_id>', methods=['GET', 'POST'])
def prepare_course_design(analysis_id):
    """
    Prepare for course design generation by collecting additional requirements.
    """
    
    current_app.logger.info("Accessed prepare_course_design for ID: %s", analysis_id)
    # Load analysis data from file
    analysis_data = load_analysis(analysis_id)
    
    if not analysis_data:
        current_app.logger.warning("Analysis ID %s not found. Redirecting to index.", analysis_id)
        flash('Analysis not found. Please submit the form to generate a new analysis.')
        return redirect(url_for('main.index'))
    
    # Check if task analysis exists
    if 'task_analysis' not in analysis_data:
        current_app.logger.warning("Task analysis not found for ID: %s. Redirecting to audience analysis.", analysis_id)      
        flash('Task analysis not found. Please complete task analysis first.')
        return redirect(url_for('main.audience_analysis', analysis_id=analysis_id))
    
    # Initialize form for additional requirements
    form = CourseDesignRequirementsForm()
    
    if form.validate_on_submit():
        current_app.logger.debug("Form submitted for ID: %s", analysis_id)
        # Get form data
        course_duration = form.course_duration.data
        delivery_format = form.delivery_format.data
        module_count = form.module_count.data
        additional_requirements = form.additional_requirements.data

        current_app.logger.debug(
            "Collected form data: duration=%s, format=%s, modules=%s",
            course_duration, delivery_format, module_count
        )
        
        # Get component options
        generate_structure = form.generate_structure.data
        generate_strategies = form.generate_strategies.data
        generate_assessment = form.generate_assessment.data
        
        # Build component list
        components = []
        if generate_structure:
            components.append('structure')
        if generate_strategies:
            components.append('strategies')
        if generate_assessment:
            components.append('assessment')
        
        current_app.logger.debug("Selected components: %s", components)

        # Make sure at least one component is selected
        if not components:
            current_app.logger.info("No components selected for ID: %s", analysis_id) 
            flash('Please select at least one component to generate.')
            return render_template('prepare_course_design.html',
                                 form=form,
                                 analysis_id=analysis_id,
                                 course_topic=analysis_data['course_topic'],
                                 audience_type=analysis_data['audience_type'],
                                 audience_analysis_html=markdown.markdown(analysis_data['audience_analysis']),
                                 task_analysis_html=markdown.markdown(analysis_data['task_analysis']),
                                 terminal_objectives=analysis_data.get('terminal_objectives', ''),
                                 current_date=analysis_data['generated_date'])
        
        # Add to analysis data
        analysis_data['course_duration'] = course_duration
        analysis_data['delivery_format'] = delivery_format
        analysis_data['module_count'] = module_count
        analysis_data['additional_requirements'] = additional_requirements
        analysis_data['design_components'] = components
        
        # Save updated analysis
        save_analysis(analysis_id, analysis_data)
        current_app.logger.info("Saved course design prep data for ID: %s", analysis_id)
        
        # Redirect to generate course design
        return redirect(url_for('main.generate_course_design', analysis_id=analysis_id))
    
    current_app.logger.debug("Rendering course design form for ID: %s", analysis_id)

    # Convert Markdown to HTML for display
    audience_analysis_html = markdown.markdown(analysis_data['audience_analysis'])
    task_analysis_html = markdown.markdown(analysis_data['task_analysis'])
    
    return render_template('prepare_course_design.html',
                          form=form,
                          analysis_id=analysis_id,
                          course_topic=analysis_data['course_topic'],
                          audience_type=analysis_data['audience_type'],
                          audience_analysis_html=audience_analysis_html,
                          task_analysis_html=task_analysis_html,
                          terminal_objectives=analysis_data.get('terminal_objectives', ''),
                          current_date=analysis_data['generated_date'])

@main.route('/generate_course_design/<analysis_id>', methods=['GET'])
def generate_course_design(analysis_id):

    current_app.logger.info("Accessed generate_course_design for ID: %s", analysis_id)
    # Load analysis data from file
    analysis_data = load_analysis(analysis_id)
    
    if not analysis_data:
        current_app.logger.warning("Analysis ID %s not found. Redirecting to index.", analysis_id)
        flash('Analysis not found. Please submit the form to generate a new analysis.')
        return redirect(url_for('main.index'))
    
    # Check if we need to regenerate any components
    components_to_generate = analysis_data.get('design_components', ['structure', 'strategies', 'assessment'])
    current_app.logger.debug("Initial requested components for generation: %s", components_to_generate)

    # Check which components already exist
    if 'course_structure' in analysis_data and 'structure' in components_to_generate:
        components_to_generate.remove('structure')
    if 'instructional_strategies' in analysis_data and 'strategies' in components_to_generate:
        components_to_generate.remove('strategies')
    if 'assessment_plan' in analysis_data and 'assessment' in components_to_generate:
        components_to_generate.remove('assessment')
    
    current_app.logger.debug("Remaining components to generate for ID %s: %s", analysis_id, components_to_generate)


    # If all components exist and no new ones are requested, redirect to view
    if not components_to_generate:
        current_app.logger.info("All components already exist for ID: %s. Redirecting to view.", analysis_id)       
        flash('Course design already exists. Redirecting to existing design.')
        return redirect(url_for('main.view_course_design', analysis_id=analysis_id))
    
    # Generate course design components
    try:
        # Extract module count if specified
        module_count = analysis_data.get('module_count')
        if module_count and isinstance(module_count, str):
            try:
                module_count = int(module_count)
            except ValueError:
                current_app.logger.warning("Invalid module_count format for ID %s: %s", analysis_id, module_count)
                module_count = None
        
        current_app.logger.info("Generating course design for ID: %s with components: %s", analysis_id, components_to_generate)

        # Generate only the requested components
        updated_analysis_data = generate_comprehensive_course_design(
            analysis_data, 
            components=components_to_generate,
            module_count=module_count
        )
        
        # Save the updated analysis
        save_analysis(analysis_id, updated_analysis_data)
        current_app.logger.info("Successfully saved generated design for ID: %s", analysis_id)

        
        # Create success message based on generated components
        component_names = {
            'structure': 'Course Structure',
            'strategies': 'Instructional Strategies',
            'assessment': 'Assessment Plan'
        }
        generated_names = [component_names[comp] for comp in components_to_generate]
        
        if len(generated_names) == 1:
            message = f"{generated_names[0]} generated successfully!"
        elif len(generated_names) == 2:
            message = f"{generated_names[0]} and {generated_names[1]} generated successfully!"
        elif len(generated_names) > 2:
            message = f"{', '.join(generated_names[:-1])} and {generated_names[-1]} generated successfully!"
        else:
            message = "Course design components generated successfully!"
            
        flash(message)
        return redirect(url_for('main.view_course_design', analysis_id=analysis_id))
    
    except Exception as e:
        current_app.logger.error("Error generating course design for ID %s: %s", analysis_id, str(e), exc_info=True)
        flash(f'Error generating course design: {str(e)}')
        return redirect(url_for('main.task_analysis', analysis_id=analysis_id))

@main.route('/view_course_design/<analysis_id>')
def view_course_design(analysis_id):

    current_app.logger.info("Accessed view_course_design for ID: %s", analysis_id)

    # Load analysis data from file
    analysis_data = load_analysis(analysis_id)
    
    if not analysis_data:
        current_app.logger.warning("Analysis ID %s not found. Redirecting to index.", analysis_id)       
        flash('Analysis not found. Please submit the form to generate a new analysis.')
        return redirect(url_for('main.index'))
    
    # Check if course design exists
    if 'course_structure' not in analysis_data:
        current_app.logger.warning("Course structure not found for ID: %s. Redirecting to generate.", analysis_id)
        flash('Course design not found. Generating now...')
        return redirect(url_for('main.generate_course_design', analysis_id=analysis_id))
    
    # Convert Markdown to HTML
    course_structure_html = markdown.markdown(analysis_data['course_structure'])
    instructional_strategies_html = markdown.markdown(analysis_data['instructional_strategies'])
    assessment_plan_html = markdown.markdown(analysis_data['assessment_plan'])
    
    return render_template('course_design.html', 
                          course_structure=course_structure_html,
                          instructional_strategies=instructional_strategies_html,
                          assessment_plan=assessment_plan_html,
                          course_topic=analysis_data['course_topic'],
                          current_date=analysis_data.get('course_design_generated_date', 
                                                      analysis_data.get('generated_date')),
                          analysis_id=analysis_id)

@main.route('/edit_course_design/<analysis_id>', methods=['GET', 'POST'])
def edit_course_design(analysis_id):
    
    current_app.logger.info("Accessed edit_course_design view for ID: %s", analysis_id)
    # Load analysis data from file
    analysis_data = load_analysis(analysis_id)
    
    if not analysis_data:
        current_app.logger.warning("Analysis ID %s not found. Redirecting to index.", analysis_id)
        flash('Analysis not found.')
        return redirect(url_for('main.index'))
    
    # Check if course design exists
    if 'course_structure' not in analysis_data:
        current_app.logger.warning("Course structure missing for ID: %s. Redirecting to generation.", analysis_id)
        flash('Course design not found. Generating now...')
        return redirect(url_for('main.generate_course_design', analysis_id=analysis_id))
    
    if request.method == 'POST':
        # Update the analysis data with edited content
        analysis_data['course_structure'] = request.form.get('course_structure', '')
        analysis_data['instructional_strategies'] = request.form.get('instructional_strategies', '')
        analysis_data['assessment_plan'] = request.form.get('assessment_plan', '')
        analysis_data['last_edited'] = datetime.now().strftime("%B %d, %Y at %H:%M")
        
        # Save the updated analysis
        save_analysis(analysis_id, analysis_data)
        current_app.logger.info("Course design updated for ID: %s", analysis_id)
        flash('Course design updated successfully!')
        
        return redirect(url_for('main.view_course_design', analysis_id=analysis_id))
    
    # For GET request, display the edit form
    return render_template('edit_course_design.html', 
                          analysis_id=analysis_id,
                          course_structure=analysis_data['course_structure'],
                          instructional_strategies=analysis_data['instructional_strategies'],
                          assessment_plan=analysis_data['assessment_plan'],
                          course_topic=analysis_data['course_topic'])

@main.route('/download_course_design/<analysis_id>')
def download_course_design(analysis_id):
   
    current_app.logger.info("Requested course design download for ID: %s", analysis_id)

    # Load analysis data from file
    analysis_data = load_analysis(analysis_id)
    
    if not analysis_data:
        current_app.logger.warning("Analysis ID %s not found. Redirecting to index.", analysis_id)     
        flash('Analysis not found.')
        return redirect(url_for('main.index'))
    
    # Check if course design exists
    if 'course_structure' not in analysis_data:
        current_app.logger.warning("Course structure not found for ID: %s. Redirecting to generation.", analysis_id)
        flash('Course design not found. Generating now...')
        return redirect(url_for('main.generate_course_design', analysis_id=analysis_id))
    
    # Create a new Word document
    doc = Document()
    
    # Add title
    title = doc.add_heading(f"Course Design: {analysis_data['course_topic']}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add generated date
    date_paragraph = doc.add_paragraph()
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_paragraph.add_run(f"Generated on {analysis_data.get('course_design_generated_date', analysis_data['generated_date'])}").italic = True
    
    doc.add_paragraph()  # Add some space
    
    # Add Course Structure
    doc.add_heading("Course Structure", 1)
    markdown_to_docx(doc, analysis_data['course_structure'])
    
    doc.add_page_break()
    
    # Add Instructional Strategies
    doc.add_heading("Instructional Strategies", 1)
    markdown_to_docx(doc, analysis_data['instructional_strategies'])
    
    doc.add_page_break()
    
    # Add Assessment Plan
    doc.add_heading("Assessment Plan", 1)
    markdown_to_docx(doc, analysis_data['assessment_plan'])
    
    # Save to a BytesIO object
    f = io.BytesIO()
    doc.save(f)
    f.seek(0)

    
    
    # Send the file for download
    return send_file(
        f,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=f"Course_Design_{analysis_data['course_topic'].replace(' ', '_')}.docx"
    )

@main.route('/results/<analysis_id>')
def results(analysis_id):

    current_app.logger.info("Accessed results view for ID: %s", analysis_id)
    # Load analysis data from file
    analysis_data = load_analysis(analysis_id)
    
    if not analysis_data:
        current_app.logger.warning("Analysis ID %s not found. Redirecting to index.", analysis_id)
        flash('Analysis not found. Please submit the form to generate a new analysis.')
        return redirect(url_for('main.index'))
    
    # Check what components exist
    has_audience = 'audience_analysis' in analysis_data
    has_task = 'task_analysis' in analysis_data
    has_course_design = 'course_structure' in analysis_data
    
    # Convert Markdown to HTML if components exist
    audience_analysis_html = markdown.markdown(analysis_data['audience_analysis']) if has_audience else None
    task_analysis_html = markdown.markdown(analysis_data['task_analysis']) if has_task else None
    
    if has_course_design:
        course_structure_html = markdown.markdown(analysis_data['course_structure'])
        instructional_strategies_html = markdown.markdown(analysis_data['instructional_strategies'])
        assessment_plan_html = markdown.markdown(analysis_data['assessment_plan'])
    else:
        course_structure_html = None
        instructional_strategies_html = None
        assessment_plan_html = None
    
    return render_template('results.html', 
                         course_topic=analysis_data['course_topic'],
                         audience_type=analysis_data['audience_type'],
                         current_date=analysis_data['generated_date'],
                         analysis_id=analysis_id,
                         audience_analysis_html=audience_analysis_html,
                         task_analysis_html=task_analysis_html,
                         course_structure_html=course_structure_html,
                         instructional_strategies_html=instructional_strategies_html,
                         assessment_plan_html=assessment_plan_html,
                         has_audience=has_audience,
                         has_task=has_task,
                         has_course_design=has_course_design)

# For backward compatibility
@main.route('/results')
def results_redirect():
    analysis_id = session.get('current_analysis_id')
    if analysis_id:
        current_app.logger.info("Redirecting to results page for ID: %s", analysis_id)
        return redirect(url_for('main.results', analysis_id=analysis_id))
    else:
        current_app.logger.warning("No analysis ID found in session. Redirecting to index.")
        flash('No analysis results found. Please submit the form first.')
        return redirect(url_for('main.index'))

@main.route('/edit_audience/<analysis_id>', methods=['GET', 'POST'])
def edit_audience_analysis(analysis_id):
    try:
        # Load analysis data from file
        analysis_data = load_analysis(analysis_id)
        
        if not analysis_data:
            current_app.logger.warning("Analysis ID %s not found. Redirecting to index.", analysis_id)
            flash('Analysis not found.')
            return redirect(url_for('main.index'))
        
        if request.method == 'POST':
            # Update only the audience analysis
            analysis_data['audience_analysis'] = request.form.get('audience_analysis', '')
            
            # Save the updated analysis
            save_analysis(analysis_id, analysis_data)
            current_app.logger.info("Audience analysis updated for ID: %s", analysis_id)
            flash('Audience analysis updated successfully!')
            
            # Redirect back to audience analysis page
            return redirect(url_for('main.audience_analysis', analysis_id=analysis_id))
        
        current_app.logger.debug("Rendering audience analysis edit form for ID: %s", analysis_id)
        # For GET request, display the edit form for audience analysis only
        return render_template('edit_audience.html', 
                            analysis_id=analysis_id,
                            audience_analysis=analysis_data['audience_analysis'],
                            course_topic=analysis_data['course_topic'])
        
    except Exception as e:
        current_app.logger.error("Error editing audience analysis for ID %s: %s", analysis_id, str(e), exc_info=True)
        flash(f"Error: {str(e)}")
        return redirect(url_for('main.index'))

@main.route('/edit_task/<analysis_id>', methods=['GET', 'POST'])
def edit_task_analysis(analysis_id):
    
    current_app.logger.info("Accessed edit_task_analysis view for ID: %s", analysis_id)

    try:
        # Load analysis data from file
        analysis_data = load_analysis(analysis_id)
        
        if not analysis_data:
            current_app.logger.warning("Analysis ID %s not found. Redirecting to index.", analysis_id)
            flash('Analysis not found.')
            return redirect(url_for('main.index'))
        
        if request.method == 'POST':
            # Update only the task analysis
            analysis_data['task_analysis'] = request.form.get('task_analysis', '')
            
            # Save the updated analysis
            save_analysis(analysis_id, analysis_data)
            current_app.logger.info("Task analysis updated for ID: %s", analysis_id)
            flash('Task analysis updated successfully!')
            
            # Redirect back to task analysis page
            return redirect(url_for('main.task_analysis', analysis_id=analysis_id))
        
        current_app.logger.debug("Rendering task analysis edit form for ID: %s", analysis_id)
        # For GET request, display the edit form for task analysis only
        return render_template('edit_task.html', 
                            analysis_id=analysis_id,
                            task_analysis=analysis_data['task_analysis'],
                            course_topic=analysis_data['course_topic'])
        
    except Exception as e:
        current_app.logger.error("Error editing task analysis for ID %s: %s", analysis_id, str(e), exc_info=True)
        flash(f"Error: {str(e)}")
        return redirect(url_for('main.index'))

@main.route('/download_audience/<analysis_id>')
def download_audience_analysis(analysis_id):

    current_app.logger.info("Download requested for audience analysis ID: %s", analysis_id)

    # Load analysis data from file
    analysis_data = load_analysis(analysis_id)
    
    if not analysis_data:
        current_app.logger.warning("Audience analysis not found for ID: %s", analysis_id)
        flash('Analysis not found.')
        return redirect(url_for('main.index'))
    
    # Create a new Word document
    doc = Document()
    
    # Add title
    title = doc.add_heading(f"Audience Analysis: {analysis_data['course_topic']}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add generated date
    date_paragraph = doc.add_paragraph()
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_paragraph.add_run(f"Generated on {analysis_data['generated_date']}").italic = True
    
    doc.add_paragraph()  # Add some space
    
    # Convert markdown to formatted Word document
    markdown_to_docx(doc, analysis_data['audience_analysis'])
    
    # Save to a BytesIO object
    f = io.BytesIO()
    doc.save(f)
    f.seek(0)
    
    # Send the file for download
    return send_file(
        f,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=f"Audience_Analysis_{analysis_data['course_topic'].replace(' ', '_')}.docx"
    )

@main.route('/download_task/<analysis_id>')
def download_task_analysis(analysis_id):

    current_app.logger.info("Download requested for task analysis ID: %s", analysis_id)

    # Load analysis data from file
    analysis_data = load_analysis(analysis_id)
    
    if not analysis_data:
        current_app.logger.warning("Audience analysis not found for ID: %s", analysis_id)
        flash('Analysis not found.')
        return redirect(url_for('main.index'))
    
    # Create a new Word document
    doc = Document()
    
    # Add title
    title = doc.add_heading(f"Task Analysis: {analysis_data['course_topic']}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add generated date
    date_paragraph = doc.add_paragraph()
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_paragraph.add_run(f"Generated on {analysis_data['generated_date']}").italic = True
    
    doc.add_paragraph()  # Add some space
    
    # Convert markdown to formatted Word document
    markdown_to_docx(doc, analysis_data['task_analysis'])
    
    # Save to a BytesIO object
    f = io.BytesIO()
    doc.save(f)
    f.seek(0)
    
    # Send the file for download
    return send_file(
        f,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=f"Task_Analysis_{analysis_data['course_topic'].replace(' ', '_')}.docx"
    )

@main.route('/generate_additional_components/<analysis_id>', methods=['POST'])
def generate_additional_components(analysis_id):

    current_app.logger.info("Initiated additional component generation for ID: %s", analysis_id)

    # Load analysis data
    analysis_data = load_analysis(analysis_id)
    
    if not analysis_data:
        current_app.logger.warning("Analysis ID %s not found.", analysis_id)
        flash('Analysis not found.')
        return redirect(url_for('main.index'))
    
    # Get components to generate
    components = []
    if request.form.get('generate_structure') and 'course_structure' not in analysis_data:
        components.append('structure')
    if request.form.get('generate_strategies') and 'instructional_strategies' not in analysis_data:
        components.append('strategies')
    if request.form.get('generate_assessment') and 'assessment_plan' not in analysis_data:
        components.append('assessment')
    
    # Generate the requested components
    if components:
        try:
            # Extract module count if specified
            module_count = analysis_data.get('module_count')
            if module_count and isinstance(module_count, str):
                try:
                    module_count = int(module_count)
                except ValueError:
                    module_count = None
                    current_app.logger.warning("Invalid module_count format in analysis ID %s", analysis_id)

            
            # Generate components
            updated_analysis_data = generate_comprehensive_course_design(
                analysis_data, 
                components=components,
                module_count=module_count
            )
            
            # Save the updated analysis
            save_analysis(analysis_id, updated_analysis_data)
            current_app.logger.info("Successfully generated components %s for ID: %s", components, analysis_id)

            
            # Create success message
            component_names = {
                'structure': 'Course Structure',
                'strategies': 'Instructional Strategies',
                'assessment': 'Assessment Plan'
            }
            generated_names = [component_names[comp] for comp in components]
            
            if len(generated_names) == 1:
                message = f"{generated_names[0]} generated successfully!"
            elif len(generated_names) == 2:
                message = f"{generated_names[0]} and {generated_names[1]} generated successfully!"
            else:
                message = f"{', '.join(generated_names[:-1])}, and {generated_names[-1]} generated successfully!"
                
            flash(message)
        except Exception as e:
            current_app.logger.error("Error generating additional components for ID %s: %s", analysis_id, str(e), exc_info=True)
            flash(f'Error generating additional components: {str(e)}')
    else:
        current_app.logger.info("No new components selected or all already exist for ID: %s", analysis_id)
        flash('No components selected for generation.')
    
    return redirect(url_for('main.view_course_design', analysis_id=analysis_id))
# Add these imports at the top of your routes.py if not already present
from models.course_materials import CourseMaterialsGenerator, generate_course_materials
import zipfile
import logging
import io
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

@main.route('/prepare_materials/<analysis_id>')
def prepare_materials(analysis_id):
    """Display the material generation preparation page."""
    current_app.logger.info("Accessed prepare_materials view for ID: %s", analysis_id)
    # Load analysis data
    analysis_data = load_analysis(analysis_id)
    
    if not analysis_data:
        current_app.logger.warning("Analysis ID %s not found.", analysis_id)
        flash('Analysis not found.')
        return redirect(url_for('main.index'))
    
    # Check if course design exists
    if 'course_structure' not in analysis_data:
        current_app.logger.warning("Course design missing for ID: %s. Redirecting to course design prep.", analysis_id)
        flash('Course design not found. Please complete course design first.')
        return redirect(url_for('main.prepare_course_design', analysis_id=analysis_id))
    
    # Extract modules from course structure
    generator = CourseMaterialsGenerator(analysis_data)
    modules = generator.modules
    
    # Check if materials already exist
    existing_materials = analysis_data.get('course_materials', {})
    if existing_materials:
            current_app.logger.info("Existing course materials found for ID: %s", analysis_id)
    else:
            current_app.logger.info("No existing course materials for ID: %s", analysis_id)

    
    return render_template('prepare_materials.html',
                          analysis_id=analysis_id,
                          course_topic=analysis_data['course_topic'],
                          audience_type=analysis_data['audience_type'],
                          design_date=analysis_data.get('course_design_generated_date', 
                                                      analysis_data['generated_date']),
                          modules=modules,
                          existing_materials=existing_materials)

@main.route('/generate_materials/<analysis_id>', methods=['POST'])
def generate_materials(analysis_id):
    """Generate course materials based on user selections including tone"""
    current_app.logger.info("Material generation initiated for ID: %s", analysis_id)

    # Load analysis data
    analysis_data = load_analysis(analysis_id)
    
    if not analysis_data:
        current_app.logger.warning("Analysis not found for ID: %s", analysis_id)
        flash('Analysis not found.')
        return redirect(url_for('main.index'))
    
    try:
        # Get form data
        selected_modules = request.form.getlist('selected_modules')
        selected_modules = [int(m) for m in selected_modules]
        
        components = request.form.getlist('components')
        detail_level = request.form.get('detail_level', 'comprehensive')
        format_preference = request.form.get('format_preference', 'structured')
        additional_notes = request.form.get('additional_notes', '')
        
        # NEW: Get tone selection
        content_tone = request.form.get('content_tone', 'default')

        current_app.logger.debug("Received form data for ID %s: modules=%s, components=%s, detail=%s, format=%s",
                                 analysis_id, selected_modules, components, detail_level, format_preference)

        
        # Validate selections
        if not selected_modules:
            current_app.logger.info("No modules selected for material generation (ID: %s)", analysis_id)
            flash('Please select at least one module.')
            return redirect(url_for('main.prepare_materials', analysis_id=analysis_id))
        
        if not components:
            current_app.logger.info("No components selected for material generation (ID: %s)", analysis_id)
            flash('Please select at least one component type.')
            return redirect(url_for('main.prepare_materials', analysis_id=analysis_id))
        
        # Validate tone selection
        valid_tones = ['default', 'optimistic', 'entertaining', 'humanized']
        if content_tone not in valid_tones:
            content_tone = 'default'                            
        # selected_modules = [int(m) for m in selected_modules]
        # Generate materials
        logger.info(f"Generating materials for modules {selected_modules} with components {components} in {content_tone} tone")
        current_app.logger.info("Generating materials for ID %s: modules=%s, components=%s",
                                analysis_id, selected_modules, components)

        
        materials = generate_course_materials(
            analysis_data,
            selected_modules=selected_modules,
            components=components,
            content_tone=content_tone,  # NEW: Pass tone parameter
            additional_notes=additional_notes
        )
        
        # Add metadata
        materials['metadata'] = {
            'generated_date': datetime.now().strftime("%B %d, %Y at %H:%M"),
            'total_modules': len(materials['modules']),
            'components_generated': components,
            'detail_level': detail_level,
            'format_preference': format_preference,
            'content_tone': content_tone,
            'additional_notes': additional_notes
        }
        
        # Merge with existing materials if any
        if 'course_materials' in analysis_data:
            
            current_app.logger.debug("Merging with existing course materials for ID: %s", analysis_id)
           
            existing_modules = {m['number']: m for m in analysis_data['course_materials'].get('modules', [])}
            
            # Update with new materials
            for new_module in materials['modules']:
                module_num = new_module['number']
                if module_num in existing_modules:
                    # Merge components
                    existing_modules[module_num]['components'].update(new_module['components'])
                else:
                    existing_modules[module_num] = new_module
            
            # Convert back to list
            materials['modules'] = list(existing_modules.values())
            materials['modules'].sort(key=lambda x: x['number'])
        
        # Save materials (use consistent key name)
        analysis_data['course_materials'] = materials
        analysis_data['materials_generated_date'] = datetime.now().strftime("%B %d, %Y at %H:%M")
        save_analysis(analysis_id, analysis_data)
        current_app.logger.info("Successfully generated materials for ID %s: modules=%d, components=%d",
                                analysis_id, len(selected_modules), len(components))

        flash(f'Successfully generated {len(components)} component(s) for {len(selected_modules)} module(s) in {content_tone} tone!')
        return redirect(url_for('main.view_materials', analysis_id=analysis_id))
        
    except Exception as e:
        current_app.logger.error("Error generating materials for ID %s: %s", analysis_id, str(e), exc_info=True)
        logger.error(f"Error generating materials: {str(e)}")
        error_message = str(e)
        flash(f'Error generating materials: {error_message}')
        return redirect(url_for('main.prepare_materials', analysis_id=analysis_id))

# Import the image service at the top of routes.py
from models.image_service import image_service

# Update the view_material route to include images
@main.route('/view_material/<analysis_id>/<int:module_id>/<material_type>')
def view_material(analysis_id, module_id, material_type):
    """View a specific material component."""
    # Load analysis data
    analysis_data = load_analysis(analysis_id)
   
    if not analysis_data or 'course_materials' not in analysis_data:
        flash('Materials not found.')
        return redirect(url_for('main.index'))
   
    # Find the module
    module = None
    for m in analysis_data['course_materials']['modules']:
        if m['number'] == module_id:
            module = m
            break
   
    if not module:
        flash('Module not found.')
        return redirect(url_for('main.view_materials', analysis_id=analysis_id))
   
    # Get the specific material
    material = module['components'].get(material_type)
   
    if not material:
        flash(f'{material_type.replace("_", " ").title()} not found for this module.')
        return redirect(url_for('main.view_materials', analysis_id=analysis_id))
   
    return render_template('view_material.html',
                          analysis_id=analysis_id,
                          course_topic=analysis_data['course_topic'],
                          module_number=module_id,
                          material_type=material_type,
                          material=material)

# Add a new route for fetching images dynamically
@main.route('/get_images/<analysis_id>/<int:module_id>')
def get_images_for_module(analysis_id, module_id):
    """Get images for a specific module via AJAX."""
    try:
        # Load analysis data
        analysis_data = load_analysis(analysis_id)
        
        if not analysis_data or 'course_materials' not in analysis_data:
            return {'error': 'Materials not found'}, 404
        
        # Find the module
        module = None
        for m in analysis_data['course_materials']['modules']:
            if m['number'] == module_id:
                module = m
                break
        
        if not module:
            return {'error': 'Module not found'}, 404
        
        # Get images
        images = image_service.get_module_images(module, count=6)
        
        return {
            'success': True,
            'images': images,
            'module_title': module.get('title', f'Module {module_id}')
        }
        
    except Exception as e:
        logger.error(f"Error fetching images for module {module_id}: {str(e)}")
        return {'error': 'Failed to fetch images'}, 500

# Update the materials dashboard to show image counts
@main.route('/view_materials/<analysis_id>')
def view_materials(analysis_id):
    """Display the materials dashboard with image information."""
    # Load analysis data
    analysis_data = load_analysis(analysis_id)
    
    if not analysis_data:
        flash('Analysis not found.')
        return redirect(url_for('main.index'))
    
    # Check if materials exist (use consistent key name)
    if 'course_materials' not in analysis_data:
        flash('No materials generated yet.')
        return redirect(url_for('main.prepare_materials', analysis_id=analysis_id))
    
    materials = analysis_data['course_materials']
    
    # Calculate statistics
    stats = calculate_materials_stats(materials)
    
    # Calculate individual stats for template
    total_lesson_plans = sum(1 for module in materials.get('modules', []) 
                           if 'lesson_plan' in module.get('components', {}))
    total_activities = sum(1 for module in materials.get('modules', []) 
                         if 'activities' in module.get('components', {}))
    total_assessments = sum(1 for module in materials.get('modules', []) 
                          if 'assessments' in module.get('components', {}))
    
     # Check image service availability
    image_service_available = any([
        image_service.unsplash_key,
        image_service.pixabay_key,
        image_service.pexels_key
    ])
    return render_template('materials_dashboard.html',
                          analysis_id=analysis_id,
                          course_topic=analysis_data['course_topic'],
                          audience_type=analysis_data['audience_type'],
                          generation_date=materials.get('metadata', {}).get('generated_date', 'Unknown'),
                          materials=materials,
                          total_lesson_plans=total_lesson_plans,
                          total_activities=total_activities,
                          total_assessments=total_assessments,
                          stats=stats,image_service_available=image_service_available)

# UPDATE the existing view_material route:
@main.route('/view_material/<analysis_id>/<int:module_id>/<material_type>')
def view_material(analysis_id, module_id, material_type):
    """View a specific material component."""
    # Load analysis data
    current_app.logger.info("Accessed materials dashboard for ID: %s", analysis_id)                                          
    analysis_data = load_analysis(analysis_id)
    
    if not analysis_data or 'course_materials' not in analysis_data:
        current_app.logger.warning("Materials not found for ID: %s", analysis_id)                                                                                    
        flash('Materials not found.')
        return redirect(url_for('main.index'))
    
    # Find the module
    module = None
    for m in analysis_data['course_materials']['modules']:
        if m['number'] == module_id:
            module = m
            break
    
    if not module:
        current_app.logger.warning("Module not found:%s", module_id)
        flash('Module not found.')
        return redirect(url_for('main.view_materials', analysis_id=analysis_id))
    
    # Get the specific material
    material = module['components'].get(material_type)
    
    if not material:
        current_app.logger.warning("Material not found for this module:%s",module_id)                                                             
        flash(f'{material_type.replace("_", " ").title()} not found for this module.')
        return redirect(url_for('main.view_materials', analysis_id=analysis_id))
    
    return render_template('view_material.html',
                          analysis_id=analysis_id,
                          course_topic=analysis_data['course_topic'],
                          module_number=module_id,
                          material_type=material_type,
                          material=material)
# UPDATE the existing edit_material route:
@main.route('/edit_material/<analysis_id>/<int:module_id>/<material_type>', methods=['GET', 'POST'])
def edit_material(analysis_id, module_id, material_type):
    """Edit a specific material component."""
    current_app.logger.info("Accessing edit view for material '%s' in module %s (analysis: %s)",
                            material_type, module_id, analysis_id)
    # Load analysis data
    analysis_data = load_analysis(analysis_id)
    
    if not analysis_data or 'course_materials' not in analysis_data:
        current_app.logger.warning("Materials not found for analysis ID: %s", analysis_id)
        flash('Materials not found.')
        return redirect(url_for('main.index'))
    
    # Find the module
    module_idx = None
    module = None
    for idx, m in enumerate(analysis_data['course_materials']['modules']):
        if m['number'] == module_id:
            module_idx = idx
            module = m
            break
    
    if not module:
        flash('Module not found.')
        return redirect(url_for('main.view_materials', analysis_id=analysis_id))
    
    # Get the specific material
    material = module['components'].get(material_type)
    
    if not material:
        current_app.logger.warning("Material type '%s' not found in module %s (analysis ID: %s)",
                                   material_type, module_id, analysis_id)
        flash(f'{material_type.replace("_", " ").title()} not found for this module.')
        return redirect(url_for('main.view_materials', analysis_id=analysis_id))
    
    if request.method == 'POST':
        # Update the material with form data
        updated_content = request.form.get('material_content', '')
        
        try:
            # Save the updated material as raw content for now
            if isinstance(material, dict):
                material['raw_content'] = updated_content
            else:
                material = {'raw_content': updated_content}
            
            analysis_data['course_materials']['modules'][module_idx]['components'][material_type] = material
            save_analysis(analysis_id, analysis_data)
            current_app.logger.info("Successfully updated material '%s' in module %s (analysis ID: %s)",
                                    material_type, module_id, analysis_id)
            
            flash(f'{material_type.replace("_", " ").title()} updated successfully!')
            return redirect(url_for('main.view_material', 
                                  analysis_id=analysis_id, 
                                  module_id=module_id, 
                                  material_type=material_type))
            
        except Exception as e:
            current_app.logger.error("Error updating material '%s' in module %s: %s",
                                     material_type, module_id, str(e), exc_info=True)
            flash(f'Error updating material: {str(e)}')
            return render_template('edit_material.html',
                                 analysis_id=analysis_id,
                                 module_id=module_id,
                                 material_type=material_type,
                                 material_content=updated_content)
    
    # For GET request, show edit form
    if isinstance(material, dict):
        material_content = material.get('raw_content', json.dumps(material, indent=2))
        
        # material_content_format = escape_control_characters(material_content)
        # material_raw = json.loads(material_content_format)
        # raw_main_content = material_raw['main_content']
        # print(raw_main_content)
    else:
        material_content = str(material)

    current_app.logger.debug("Loaded material '%s' for editing (module: %s, analysis ID: %s)",
                             material_type, module_id, analysis_id)
    
    return render_template('edit_material.html',
                          analysis_id=analysis_id,
                          module_id=module_id,
                          material_type=material_type,
                          material_content=material_content)

@main.route('/generate_single_material/<analysis_id>/<int:module_id>/<material_type>')
def generate_single_material(analysis_id, module_id, material_type):
    """Generate a single material component for a specific module."""
    current_app.logger.info(
        "Generating single material '%s' for module %s (analysis ID: %s)",
        material_type, module_id, analysis_id
    )
    # Load analysis data
    analysis_data = load_analysis(analysis_id)
    
    if not analysis_data:
        current_app.logger.warning("Analysis not found for ID: %s", analysis_id)
        flash('Analysis not found.')
        return redirect(url_for('main.index'))
    
    try:
        # Generate the specific material
        generator = CourseMaterialsGenerator(analysis_data)
        current_app.logger.debug("Calling material generator for component: %s", component)
        
        # Map material type to component name
        component_map = {
            'lesson_plan': 'lesson_plans',
            'content': 'content',
            'activities': 'activities',
            'assessments': 'assessments',
            'instructor_guide': 'instructor_guides'
        }
        
        component = component_map.get(material_type)
        if not component:
            flash('Invalid material type.')
            return redirect(url_for('main.view_materials', analysis_id=analysis_id))
        
        # Generate material for single module and component
        new_materials = generator.generate_all_materials(
            selected_modules=[module_id],
            components=[component]
        )
        
        # Initialize course_materials if it doesn't exist
        if 'course_materials' not in analysis_data:
            analysis_data['course_materials'] = {'modules': []}
        
        # Find and update the specific module
        module_found = False
        for idx, module in enumerate(analysis_data['course_materials']['modules']):
            if module['number'] == module_id:
                # Update the specific component
                if len(new_materials['modules']) > 0:
                    new_component = new_materials['modules'][0]['components'].get(material_type)
                    if new_component:
                        if 'components' not in module:
                            module['components'] = {}
                        module['components'][material_type] = new_component
                module_found = True
                break
        
        if not module_found and len(new_materials['modules']) > 0:
            # Module not found in existing materials, add it
            analysis_data['course_materials']['modules'].append(new_materials['modules'][0])
            analysis_data['course_materials']['modules'].sort(key=lambda x: x['number'])
        
        # Save updated analysis
        save_analysis(analysis_id, analysis_data)
        current_app.logger.info("Material '%s' generated and saved successfully for module %s", material_type, module_id)

        
        flash(f'{material_type.replace("_", " ").title()} generated successfully!')
        return redirect(url_for('main.view_material', 
                              analysis_id=analysis_id, 
                              module_id=module_id, 
                              material_type=material_type))
        
    except Exception as e:
        current_app.logger.error("Error generating single material: %s", str(e), exc_info=True)
        logger.error(f"Error generating single material: {str(e)}")
        flash(f'Error generating material: {str(e)}')
        return redirect(url_for('main.view_materials', analysis_id=analysis_id))

@main.route('/regenerate_module/<analysis_id>/<int:module_id>')
def regenerate_module(analysis_id, module_id):
    """Regenerate all materials for a specific module."""
    current_app.logger.info("Starting regeneration for module %s (analysis ID: %s)", module_id, analysis_id)

    try:
        # Load analysis data
        analysis_data = load_analysis(analysis_id)
        
        if not analysis_data:
            current_app.logger.warning("Analysis not found for ID: %s", analysis_id)
            flash('Analysis not found.')
            return redirect(url_for('main.index'))
        
        current_app.logger.debug("Loaded analysis data for regeneration")
        
        # Generate all materials for the module
        generator = CourseMaterialsGenerator(analysis_data)
        materials = generator.generate_all_materials(
            selected_modules=[module_id],
            components=['lesson_plans', 'content', 'activities', 'assessments', 'instructor_guides']
        )

        current_app.logger.debug("Generated new materials for module %s", module_id)
        
        # Initialize course_materials if it doesn't exist
        if 'course_materials' not in analysis_data:
            analysis_data['course_materials'] = {'modules': []}
            current_app.logger.info("Initialized empty course_materials in analysis data")
        
        # Replace the module materials
        modules = analysis_data['course_materials']['modules']
        # Remove existing module
        modules[:] = [mod for mod in modules if mod['number'] != module_id]
        # Add regenerated module
        if len(materials['modules']) > 0:
            modules.extend(materials['modules'])
            modules.sort(key=lambda x: x['number'])
        
        # Save updated analysis
        save_analysis(analysis_id, analysis_data)
        current_app.logger.info("Regenerated module %s successfully saved", module_id)
        
        flash(f'Module {module_id} regenerated successfully!')
        return redirect(url_for('main.view_materials', analysis_id=analysis_id))
        
    except Exception as e:
        current_app.logger.error("Error regenerating module %s: %s", module_id, str(e), exc_info=True)
        flash(f'Error regenerating module {module_id}: {str(e)}')
        return redirect(url_for('main.view_materials', analysis_id=analysis_id))



def clean_markdown_text(text):
    if not isinstance(text, str):
        return text

    text = text.encode("utf-8").decode("unicode_escape")  # Convert \n, \", etc.
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)          # Remove bold (**bold**)
    text = re.sub(r'^\s*[-*]\s+', '', text, flags=re.MULTILINE)  # Remove bullet points
    text = re.sub(r'\\n', '\n', text)                     # Clean up double-escaped \n
    text = re.sub(r'\\', '', text)                        # Remove leftover backslashes
    return text.strip()

def clean_json_structure(obj):
    if isinstance(obj, dict):
        return {k: clean_json_structure(v) for k, v in obj.items()}
    elif isinstance(obj, list):
        return [clean_json_structure(item) for item in obj]
    elif isinstance(obj, str):
        return clean_markdown_text(obj)
    else:
        return obj
import json
import re

def clean_groq_response(raw_response):
    # Step 1: Remove outer single quotes
    if raw_response.startswith("'") and raw_response.endswith("'"):
        raw_response = raw_response[1:-1]

    # Step 2: Remove Markdown-style ```json and ``` wrappers
    if raw_response.startswith("```json") or raw_response.startswith("```"):
        raw_response = re.sub(r"^```json\s*|```$", "", raw_response).strip()

    # Step 3: Decode escaped characters
    try:
        raw_response = raw_response.encode().decode('unicode_escape')
    except Exception as e:
        logger.warning(f"Unicode decode issue: {e}")

    # Step 4: Try to parse JSON
    try:
        return json.loads(raw_response)
    except json.JSONDecodeError as e:
        raise ValueError(f"Failed to decode JSON from Groq response: {e}")


@main.route('/download_all_materials/<analysis_id>')
def download_all_materials(analysis_id):
    """Download all materials as a ZIP file."""
    scorm_version = request.args.get('scorm', '').strip()
    include_scorm = scorm_version == '2004'

    # Load analysis data
    analysis_data = load_analysis(analysis_id)
    
    if not analysis_data or 'course_materials' not in analysis_data:
        current_app.logger.error("While download_all_materials-Materials not found: %s", analysis_id, str(e), exc_info=True)
        flash('Materials not found.')
        return redirect(url_for('main.view_materials', analysis_id=analysis_id))
    
    try:
        # Create a ZIP file in memory
        zip_buffer = io.BytesIO()
        
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            # Add overview document
            overview = create_materials_overview(analysis_data)
            zip_file.writestr('Course_Materials_Overview.txt', overview)
            
            navigation_data = create_combined_navigation(analysis_data['course_materials'])
            zip_file.writestr('navigation.json', json.dumps(navigation_data, indent=2))
			# Add SCORM runtime files if requested
            if include_scorm:
                scorm_files = {
                    'index_lms.html': os.path.join(current_app.root_path, 'static', 'index_lms.html'),
                    'course.js': os.path.join(current_app.root_path, 'static', 'js', 'course.js'),
                    'ADLwrapper.js': os.path.join(current_app.root_path, 'static', 'js', 'ADLwrapper.js'),
                    'index.css': os.path.join(current_app.root_path, 'static', 'css', 'index.css')                                 
                }
                
                # Add all files from the images folder
                images_folder = os.path.join(current_app.root_path, 'static', 'images')
                for root, dirs, files in os.walk(images_folder):
                    for file in files:
                        abs_path = os.path.join(root, file)
                        # Get relative path inside 'static' to preserve folder structure
                        rel_path = os.path.relpath(abs_path, os.path.join(current_app.root_path, 'static'))
                        scorm_files[rel_path] = abs_path                      
                for arcname, filepath in scorm_files.items():
                    if os.path.exists(filepath):
                        zip_file.write(filepath, arcname)
                    else:
                        current_app.logger.warning(f'SCORM file missing: {filepath}')
                
                # Generate and add manifest
                manifest_xml = create_scorm_manifest(
                    course_title=analysis_data['course_topic'],
                    modules=analysis_data['course_materials']['modules']
                )
                zip_file.writestr('imsmanifest.xml', manifest_xml)

         
            # Add each module's materials
            for module in analysis_data['course_materials']['modules']:
                module_folder = f"Module_{module['number']}_{sanitize_filename(module['title'])}"
                module_title = module['title']
                module_number = module['number']
                # Collect chapter/component names
                # chapters = []
                # for component_type, component_data in module['components'].items():
                #     if component_data:
                #         # Format component type as title (e.g., "lesson_plan" -> "Lesson Plan")
                #         chapter_name = component_type.replace('_', ' ').title()
                #         chapters.append(chapter_name)
                
                
                # Add each component
                for component_type, component_data in module['components'].items():
                    if component_data:
                        # Create filename
                        filename = f"{module_folder}/{component_type.replace('_', ' ').title()}.json"
                        
                        # Add to ZIP
                        content = json.dumps(component_data, indent=2)
                        zip_file.writestr(filename, content)
                        
                        # Also create a formatted text version
                        text_content = format_material_as_text(component_type, component_data)
                        text_filename = f"{module_folder}/{component_type.replace('_', ' ').title()}.txt"
                        zip_file.writestr(text_filename, text_content)

                       # 3. Clean JSON (structured � parsed from markdown if needed)
                        component_type = component_type.lower()

                        if component_type in ["assessments", "assessment"]:
                            print('assessments')
                            if 'raw_content' in component_data:
                                asses_json_output = parse_assessment_to_json(component_data['raw_content'])
                                cleaned_data = asses_json_output
                            elif 'comprehensive_assessments' in component_data:
                                #asses_json_output = parse_assessment_to_json(component_data['comprehensive_assessments'])
                                combined_content = (
                                    component_data.get('comprehensive_assessments', '') + '\n' +
                                    component_data.get('practice_questions', '')
                                )
                                asses_json_output = parse_assessment_to_json(combined_content)
                                cleaned_data = asses_json_output
 
                            else:
                                asses_json_output = None  # or handle the missing key case appropriately
 

                        elif component_type == "content":
                             #json_output = parse_content_to_json_contenttype(component_data['raw_content'])
                            if 'raw_content' in component_data:
                                json_output = parse_content_to_json_contenttype(component_data['raw_content'],analysis_id,component_data['metadata'])
                            elif 'main_content' in component_data:
                                json_output = parse_content_to_json_contenttype(component_data['main_content'],analysis_id,component_data['metadata'])
                            else:
                                json_output = None  # or handle the missing key case appropriately

                            cleaned_data = json_output
                                                     
                        else:
                            cleaned_data = component_data  # fallback: use raw JSON if unknown type
                            

                        clean_json_filename = f"{filename}.clean.json"
                        clean_json_content = json.dumps(cleaned_data, indent=2)
                        zip_file.writestr(clean_json_filename, clean_json_content)


            # 🔹 Now SCORM logic (AFTER modules)
            if include_scorm:
                scorm_files = {
                    'index_lms.html': os.path.join(current_app.root_path, 'static', 'index_lms.html'),
                    'course.js': os.path.join(current_app.root_path, 'static', 'js', 'course.js'),
                    'ADLwrapper.js': os.path.join(current_app.root_path, 'static', 'js', 'ADLwrapper.js'),
                    'index.css': os.path.join(current_app.root_path, 'static', 'css', 'index.css')
                }

                # Add all files from static/images
                images_folder = os.path.join(current_app.root_path, 'static', 'images')
                for root, dirs, files in os.walk(images_folder):
                    for file in files:
                        abs_path = os.path.join(root, file)
                        rel_path = os.path.relpath(abs_path, os.path.join(current_app.root_path, 'static'))
                        scorm_files[rel_path] = abs_path

                # Point to the subfolder for the current analysis_id only
                analysis_folder = os.path.join(
                    current_app.root_path, "..", "moduleimages", f"{analysis_id}"
                )
                analysis_folder = os.path.abspath(analysis_folder)

                if os.path.exists(analysis_folder):
                    for root, dirs, files in os.walk(analysis_folder):
                        for file in files:
                            abs_path = os.path.join(root, file)

                            # Make rel_path relative to the *parent of moduleimages*
                            rel_path = os.path.relpath(abs_path, os.path.join(analysis_folder, "..", ".."))
                            rel_path = rel_path.replace("\\", "/")

                            # ✅ Now SCORM will contain: moduleimages/analysis_<id>/image.jpg
                            scorm_files[rel_path] = abs_path

                for arcname, filepath in scorm_files.items():
                    if os.path.exists(filepath):
                        zip_file.write(filepath, arcname)
                    else:
                        current_app.logger.warning(f'SCORM file missing: {filepath}')

                # Add manifest
                manifest_xml = create_scorm_manifest(
                    course_title=analysis_data['course_topic'],
                    modules=analysis_data['course_materials']['modules']
                )
                zip_file.writestr('imsmanifest.xml', manifest_xml)

        # Prepare for download
        zip_buffer.seek(0)
        
        return send_file(
            zip_buffer,
            mimetype='application/zip',
            as_attachment=True,
            download_name=f"Course_Materials_{sanitize_filename(analysis_data['course_topic'])}.zip"
        )
        
    except Exception as e:
        current_app.logger.error("Error creating materials ZIP: %s", str(e), exc_info=True)
        flash('Error creating download file.')
        return redirect(url_for('main.view_materials', analysis_id=analysis_id))

import re


def format_question_sections(text):
    """
    Inserts a newline after any line that starts with a variant of 'Question <number>'
    (e.g., 'Question1', 'Question  2:', etc.), including optional bullets and punctuation.
    """
    question_pattern = r'(^[\*\+\-\s]*\**\s*Question\s*\d+\s*[:\-]?)\s*'
    text = re.sub(question_pattern, r'\1\n', text, flags=re.IGNORECASE | re.MULTILINE)
    return text         


def parse_assessment_to_json(assessment_text):
    # Define section assessmentHeadings for the assessment
    assessmentHeadings = [
        "Comprehensive Assessment Suite", "Knowledge Check Questions",
        "Multiple Choice Questions", "True/False Questions",
        "Short Answer Questions", "Application Questions",
        "Scenario-Based Questions", "Analysis and Synthesis Questions",
        "Practical Assessment Project",
        "Self-Assessment Tools","Answer Keys and Explanations","Practice Questions for"
    ]
    
    single_instance_sections = {"Multiple Choice Questions", "True/False Questions"}
    assessmentRaw_text = format_question_sections(assessment_text) 
    def asmt_clean_line(line):
        cleaned = re.sub(r'^#+\s*', '', line)
        cleaned = re.sub(r'^[-*+]\s+', '', cleaned)
        cleaned = re.sub(r'^\d+\.\s+', '', cleaned)
        cleaned = re.sub(r'\*\*|__|\*|_', '', cleaned)
        cleaned = re.sub(r'`+', '', cleaned)
        cleaned = re.sub(r':\s*$', '', cleaned)
        cleaned = re.sub(r'\t', ' ', cleaned)
        return cleaned.strip()
    
    asmt_raw_lines = [line.strip() for line in assessmentRaw_text.split('\n') if line.strip()]
    asmt_cleaned_lines = [asmt_clean_line(line) for line in asmt_raw_lines]

    asmt_parsed_data = {}
    asmt_current_heading = None
    asmt_lowercase_headings = [h.lower() for h in assessmentHeadings]
    seen_once_set = set()  # track first-seen single-instance headings

    # for line in asmt_cleaned_lines:
    #     if line.lower() in asmt_lowercase_headings:
    #         asmt_matching_heading = next(h for h in assessmentHeadings
    #                                 if h.lower() == line.lower())
    #         asmt_current_heading = asmt_matching_heading
    #         if asmt_current_heading not in asmt_parsed_data:
    #             asmt_parsed_data[asmt_current_heading] = []
    #     elif asmt_current_heading:
    #         asmt_parsed_data[asmt_current_heading].append(line)

    # for line in asmt_cleaned_lines:
    #     if any(line.lower().startswith(h.lower()) for h in assessmentHeadings if h == "Practice Questions for") or line.lower() in asmt_lowercase_headings:
    #         matching_heading = next(h for h in assessmentHeadings if line.lower().startswith(h.lower())) if "Practice Questions for" in assessmentHeadings and line.lower().startswith("practice questions for") else next(h for h in assessmentHeadings if h.lower() == line.lower())
    #         asmt_current_heading = matching_heading
    #         if asmt_current_heading not in asmt_parsed_data:
    #             asmt_parsed_data[asmt_current_heading] = []
    #     elif asmt_current_heading:
    #         asmt_parsed_data[asmt_current_heading].append(line)
    for line in asmt_cleaned_lines:
        # Check if the line starts with or contains a heading (ignoring parenthetical content)
        matching_heading = None
        for h in assessmentHeadings:
            h_lower = h.lower()
            # Remove parenthetical content for comparison
            asmt_cleaned_lines = re.sub(r'\s*\([^)]*\)', '', line.lower()).strip()
            if asmt_cleaned_lines == h_lower or (h == "Practice Questions for" and line.lower().startswith(h_lower)):
                if h in single_instance_sections and h in seen_once_set:
                    matching_heading = None  # Skip duplicate occurrence
                    break                                                                                                                   
                matching_heading = h
                seen_once_set.add(h)     
                break
        if matching_heading:
            asmt_current_heading = matching_heading
            if asmt_current_heading not in asmt_parsed_data:
                asmt_parsed_data[asmt_current_heading] = []
        elif asmt_current_heading:
            asmt_parsed_data[asmt_current_heading].append(line)
           

    asmt_standard_json = {
        "comprehensive_assessments": {
            "knowledge_check_questions": {
                "multiple_choice_questions": [],
                "true_false_questions": [],
                "short_answer_questions": []
            },
            "application_questions": {
                "scenario_based_questions": []
            },
            "analysis_and_synthesis_questions": [],
            "practical_assessment_project": {
                "project_description": "",
                "project_requirements": [],
                "deliverables": [],
                "grading_rubric": {}
            },
            "self_assessment_tools": {
                "knowledge_self_check": [],
                "skills_self_assessment": []
            },
            #"answer_keys_and_explanations": [],
            "practice_questions": []
        },
        "assessment_overview": {
            "total_questions": "",
            "question_types": [],
            "assessment_features": [],
            "estimated_assessment_time": ""
        }
    }

    def process_multiple_choice(lines):
        questions = []
        current_question = {}
        current_options = []
        expect_question_text = False
        for line in lines:
            line = asmt_clean_line(line)
            if line.startswith("Question"):
                if current_question:
                    current_question["options"] = current_options
                    questions.append(current_question)
                    current_options = []
                    current_question = {}
                
                current_question = {
                    "question_number": len(questions) + 1,
                    "question": "",
                    "options": [],
                    "correct_answer": "",
                    "content_reference": "",
                    "learning_objective_tested": ""
                }
                expect_question_text = True
            #elif expect_question_text and line and not line.startswith(("a)", "b)", "c)", "d)", "Correct Answer:", "Content Reference:", "Learning Objective Tested:")):
            elif expect_question_text:                          
                if line.strip():  # Skip blank/empty lines                     
                    current_question["question"] = line.strip()
                    expect_question_text = False
            elif line.startswith(("a)", "b)", "c)", "d)")):
                 current_options.append(line.strip())
            elif line.startswith("Correct Answer:"):
                current_question["correct_answer"] = line.split(":", 1)[1].strip()
            elif line.startswith("Content Reference:"):
                current_question["content_reference"] = line.split(":", 1)[1].strip()
            elif line.startswith("Learning Objective Tested:"):
                current_question["learning_objective_tested"] = line.split(":", 1)[1].strip()
        if current_question:
            current_question["options"] = current_options
            questions.append(current_question)
        return questions

    def process_true_false(lines):
        questions = []
        current_question = {}
        expect_question_text = False                            
        for line in lines:
            line = asmt_clean_line(line)
            if line.startswith("Question"):
                if current_question:
                    questions.append(current_question)
                    current_question = {}
                expect_question_text = True
                current_question = {
                    "question_number": len(questions) + 1,
                    "question": "",
                    "correct_answer": False,
                    "content_reference": "",
                    "learning_objective_tested": ""
                }
            elif expect_question_text:
                if line.strip():  # non-empty line just after 'Question'
                    current_question["question"] = line.strip()
                    expect_question_text = False
            elif line.startswith("Correct Answer:"):
                answer = line.split(":", 1)[1].strip()
                current_question["correct_answer"] = answer.lower().startswith(
                    "true")
            elif line.startswith("Content Reference:"):
                current_question["content_reference"] = line.split(
                    ":", 1)[1].strip()
            elif line.startswith("Learning Objective Tested:"):
                current_question["learning_objective_tested"] = line.split(
                    ":", 1)[1].strip()
        if current_question:
            questions.append(current_question)
        return questions

    def process_short_answer(lines):
        questions = []
        current_question = {}
        key_points = []
        expect_question_text = False               
        for line in lines:
            line = asmt_clean_line(line)
            if line.startswith("Question"):
                if current_question:
                    current_question["key_points_required"] = key_points
                    questions.append(current_question)
                    key_points = []
                    current_question = {}
                expect_question_text = True
                current_question = {
                    "question_number": len(questions) + 1,
                    "question": "",
                    "sample_correct_answer": "",
                    "key_points_required": [],
                    "content_reference": "",
                    "learning_objective_tested": ""
                }

            elif expect_question_text and line and not line.startswith(("Sample Correct Answer:", "Key Points Required:","Content Reference:","Learning Objective Tested:")):
                current_question["question"] = line.strip()
                expect_question_text = False    
            elif line.startswith("Sample Correct Answer:"):
                current_question["sample_correct_answer"] = line.split(
                    ":", 1)[1].strip()
            elif line.startswith("Key Points Required:"):
                key_points = line.split(":", 1)[1].strip().split(", ")
            elif line.startswith("Content Reference:"):
                current_question["content_reference"] = line.split(
                    ":", 1)[1].strip()
            elif line.startswith("Learning Objective Tested:"):
                current_question["learning_objective_tested"] = line.split(
                    ":", 1)[1].strip()
        if current_question:
            current_question["key_points_required"] = key_points
            questions.append(current_question)
        return questions

    def process_scenario_based(lines):
        questions = []
        current_question = {}
        expect_question_text = False                     
        rubric = {}
        for line in lines:
            line = asmt_clean_line(line)
            if line.startswith("Question"):
                if current_question:
                    current_question["assessment_rubric"] = rubric
                    questions.append(current_question)
                    rubric = {}
                    current_question = {}
                expect_question_text = True
                current_question = {
                    "question_number": len(questions) + 1,
                    "question": "",
                    "sample_correct_answer": "",
                    "assessment_rubric": {},
                    "content_connection": ""
                }
            elif expect_question_text and line and not line.startswith(("Question","Sample Correct Answer",
                                 "Assessment Rubric",
                                 "Excellent", "Good", "Satisfactory", "Needs Improvement","Content Connection")):
                current_question["question"] = line.strip()
                expect_question_text = False 

            elif line.startswith("Sample Correct Answer:"):
                current_question["sample_correct_answer"] = line.split(
                    ":", 1)[1].strip()
            elif line.startswith("Assessment Rubric:"):
                continue
            elif line.startswith(
                ("Excellent", "Good", "Satisfactory", "Needs Improvement")):
                score = 4 if line.startswith(
                    "Excellent") else 3 if line.startswith(
                        "Good") else 2 if line.startswith(
                            "Satisfactory") else 1
                description = line.split("(", 1)[0].strip()
                rubric[description.lower()] = {
                    "score": score,
                    "description": line.split(":", 1)[1].strip()
                }
            elif line.startswith("Content Connection:"):
                current_question["content_connection"] = line.split(
                    ":", 1)[1].strip()
        if current_question:
            current_question["assessment_rubric"] = rubric
            questions.append(current_question)
        return questions

    def process_analysis_synthesis(lines):
        questions = []
        current_question = {}
        grading_criteria = []
        expect_question_text = False                     
        for line in lines:
            line = asmt_clean_line(line)
            if line.startswith("Question"):
                if current_question:
                    current_question["grading_criteria"] = grading_criteria
                    questions.append(current_question)
                    grading_criteria = []
                    current_question = {}
                expect_question_text = True
                current_question = {
                    "question_number": len(questions) + 1,
                    "question": "",
                    "sample_answer": "",
                    "grading_criteria": [],
                    "content_references": ""
                }

            elif expect_question_text and line and not line.startswith(("Question","Sample Answer",
                                 "Grading Criteria",
                                 "Content References")):
                current_question["question"] = line.strip()
                expect_question_text = False 

            elif line.startswith("Sample Answer:"):
                current_question["sample_answer"] = line.split(":",
                                                               1)[1].strip()
            elif line.startswith("Grading Criteria:"):
                grading_criteria = line.split(":", 1)[1].strip().split(", ")
            elif line.startswith("Content References:"):
                current_question["content_references"] = line.split(
                    ":", 1)[1].strip()
        if current_question:
            current_question["grading_criteria"] = grading_criteria
            questions.append(current_question)
        return questions

    def process_practical_project(lines):
        project = {
            "project_description": "",
            "project_requirements": [],
            "deliverables": [],
            "grading_rubric": {}
        }
        current_section = None
        requirements = []
        deliverables = []
        rubric = {}
        subheadings = ["Project Description", "Project Requirements", "Deliverables", "Grading Rubric"]
        lowercase_subheadings = [h.lower() for h in subheadings]

        for line in lines:
            line = asmt_clean_line(line)
            if line.lower() in lowercase_subheadings:
                current_section = next(h for h in subheadings if h.lower() == line.lower()).lower().replace(" ", "_")
            elif current_section and line:
                if current_section == "project_description":
                    project["project_description"] = line
                elif current_section == "project_requirements":
                    requirements.append(line)
                elif current_section == "deliverables":
                    deliverables.append(line)
                elif current_section == "grading_rubric":
                    try:
                        key = line.split("(", 1)[0].strip().lower().replace(" ", "_")
                        weight = int(line.split("(", 1)[1].split("%")[0])
                        description = line.split(":", 1)[1].strip()
                        rubric[key] = {"weight": weight, "description": description}
                    except (IndexError, ValueError) as e:
                        print(f"Error processing grading rubric line '{line}': {str(e)}")

        project["project_requirements"] = requirements
        project["deliverables"] = deliverables
        project["grading_rubric"] = rubric
        return project

    def process_self_assessment(lines):
        self_assessment = {
            "knowledge_self_check": [],
            "skills_self_assessment": []
        }
        current_section = None
        subheadings = ["Knowledge Self-Check", "Skills Self-Assessment"]
        lowercase_subheadings = [h.lower() for h in subheadings]

        for line in lines:
            line = asmt_clean_line(line)
            if line.lower() in lowercase_subheadings:
                current_section = next(h for h in subheadings if h.lower() == line.lower()).lower().replace(" ", "_")
            elif current_section and line:
                if current_section == "knowledge_self-check":
                 scale = "1-5" if "(1-5)" in line else ""
                 question_text = line.split("(1-5)")[0].strip() if "(1-5)" in line else line
                 self_assessment["knowledge_self_check"].append({"question": question_text, "scale": scale})
                elif current_section == "skills_self-assessment":
                 options = ["Yes", "No", "Partially"] if "Yes/No/Partially" in line else []
                 question_text = line.split("(Yes/No/Partially)")[0].strip() if "(Yes/No/Partially)" in line else line
                 self_assessment["skills_self_assessment"].append({"question": question_text, "options": options})

        return self_assessment

    # def process_answer_keys(lines):
    #     answer_keys = []
    #     current_question = {}
    #     common_wrong_answers = []
    #     for line in lines:
    #         line = asmt_clean_line(line)
    #         if line.startswith("Question"):
    #             if current_question:
    #                 current_question[
    #                     "common_wrong_answers"] = common_wrong_answers
    #                 answer_keys.append(current_question)
    #                 common_wrong_answers = []
    #                 current_question = {}
    #             question_text = line.split(
    #                 ":", 1)[1].strip() if ":" in line else line.replace(
    #                     "Question", "").strip()
    #             current_question = {
    #                 "question": question_text,
    #                 "correct_answer": "",
    #                 "explanation": "",
    #                 "common_wrong_answers": [],
    #                 "content_reference": "",
    #                 "tips": ""
    #             }
    #         elif line.startswith("Correct Answer:"):
    #             current_question["correct_answer"] = line.split(":",
    #                                                             1)[1].strip()
    #         elif line.startswith("Explanation:"):
    #             current_question["explanation"] = line.split(":", 1)[1].strip()
    #         elif line.startswith("Content Reference:"):
    #             current_question["content_reference"] = line.split(
    #                 ":", 1)[1].strip()
    #         elif line.startswith("Common Wrong Answers:"):
    #             common_wrong_answers = [
    #                 ans.strip() for ans in line.split(":", 1)[1].split(", ")
    #             ] if ":" in line else []
    #         elif line.startswith("Tips:"):
    #             current_question["tips"] = line.split(":", 1)[1].strip()
    #     if current_question:
    #         current_question["common_wrong_answers"] = common_wrong_answers
    #         answer_keys.append(current_question)
    #     return answer_keys
    
    def process_practice_questions(lines):
        questions = []
        current_question = {}
        current_options = []
        expect_question_text = False
        for line in lines:
            line = asmt_clean_line(line)
            if line.startswith("Practice Question"):
                if current_question:
                    current_question["options"] = current_options
                    questions.append(current_question)
                    current_options = []
                    current_question = {}
                expect_question_text = True
                current_question = {
                    "question_number": len(questions) + 1,
                    "question": "",
                    "options": [],
                    "answer": "",
                    "content_reference": "",
                    "study_tip": ""
                }
            elif expect_question_text and line and not line.startswith(("A)", "B)", "C)", "D)", "Answer:", "Content Reference:", "Study Tip:")):
                current_question["question"] = line.strip()
                expect_question_text = False
            elif line.startswith(("A)", "B)", "C)", "D)")):
                current_options.append(line[2:].strip())
            elif line.startswith("Answer:"):
                current_question["answer"] = line.split(":", 1)[1].strip()
            elif line.startswith("Content Reference:"):
                current_question["content_reference"] = line.split(":", 1)[1].strip()
            elif line.startswith("Study Tip:"):
                current_question["study_tip"] = line.split(":", 1)[1].strip()
        if current_question:
            current_question["options"] = current_options
            questions.append(current_question)
        return questions

    for key, value in asmt_parsed_data.items():
        if key == "Multiple Choice Questions":
         asmt_standard_json["comprehensive_assessments"][
             "knowledge_check_questions"][
                 "multiple_choice_questions"] = process_multiple_choice(value)

        elif key == "True/False Questions":
            asmt_standard_json["comprehensive_assessments"][
                "knowledge_check_questions"][
                    "true_false_questions"] = process_true_false(value)
        elif key == "Short Answer Questions":
            asmt_standard_json["comprehensive_assessments"][
                "knowledge_check_questions"][
                    "short_answer_questions"] = process_short_answer(value)
        elif key == "Scenario-Based Questions":
            asmt_standard_json["comprehensive_assessments"][
                "application_questions"][
                    "scenario_based_questions"] = process_scenario_based(value)
        elif key == "Analysis and Synthesis Questions":
            asmt_standard_json["comprehensive_assessments"][
                "analysis_and_synthesis_questions"] = process_analysis_synthesis(
                    value)
        elif key == "Practical Assessment Project":
            asmt_standard_json["comprehensive_assessments"][
                "practical_assessment_project"] = process_practical_project(
                    value)
        elif key == "Self-Assessment Tools":
            asmt_standard_json["comprehensive_assessments"][
                "self_assessment_tools"] = process_self_assessment(value)
            
        elif key.startswith("Practice Questions for"):
            asmt_standard_json["comprehensive_assessments"]["practice_questions"] = process_practice_questions(value)
        
        # elif key == "Answer Keys and Explanations":
        #     asmt_standard_json["comprehensive_assessments"][
        #         "answer_keys_and_explanations"] = process_answer_keys(value)

    return asmt_standard_json



def create_combined_navigation(course_materials):
    navigation_data = []

    for module in course_materials["modules"]:
        module_number = module["number"]
        module_title = module["title"]
        module_folder = f"Module_{module_number}_{sanitize_filename(module_title)}"

        navigation_data.append({
            "module": module_title,
            "path": module_folder,
            "hasAssessment": True
        })

    return navigation_data

def remove_trailing_commas(json_str):
    # Remove trailing commas inside objects and arrays
    # 1. Trailing commas in objects: { "a": 1, }
    json_str = re.sub(r',(\s*})', r'\1', json_str)
    # 2. Trailing commas in arrays: [1, 2, ]
    json_str = re.sub(r',(\s*])', r'\1', json_str)
    return json_str

def parse_markdown_to_scorm_object(md: str):
    import re

    def extract_section(name, content):
        # Match a bold label and extract until the next bold label
        pattern = rf"\*\*{re.escape(name)}\*\*[:]?\s*\n?(.*?)(?=\n\s*\*\*[^*\n]+\*\*|\Z)"
        match = re.search(pattern, content, re.DOTALL)
        return match.group(1).strip() if match else ""

    def extract_core_concepts(section):
        def get_text(label):
            return extract_section(label, section)

        def get_list(label):
            pattern = rf"\*\*{re.escape(label)}\*\*[:]?\s*\n((?:\s*[-+*]\s+.*\n?)+)"
            match = re.search(pattern, section)
            return [item.strip() for item in re.findall(r"[-+*] (.+)", match.group(1))] if match else []

        return {
            "definition": get_text("Definition"),
            "theoreticalFoundation": get_text("Theoretical Foundation"),
            "keyComponents": get_list("Key Components")
        }

    def extract_best_practices(section):
        pattern = r"\*\*Best Practices\*\*\n((?:\s*[-+*]\s+.*\n?)+)"
        match = re.search(pattern, section)
        return re.findall(r"[-+*] (.*)", match.group(1)) if match else []

    def clean_title(header):
        return re.sub(r'[#:*\n]', '', header).strip()

    # Step 1: Locate the Detailed Topic Coverage section
    topic_coverage_match = re.search(r"### Detailed Topic Coverage\n+(.*?)(?=\n### |\n## |\Z)", md, re.DOTALL)
    if not topic_coverage_match:
        return {"chapter": {"topics": []}}

    topic_content = topic_coverage_match.group(1).strip()

    # Step 2: Split by all #### Subtopics inside Detailed Topic Coverage
    topic_sections = re.split(r"\n(?=#### )", topic_content)

    topics = []
    for section in topic_sections:
        title_match = re.match(r"#### (.+)", section)
        if not title_match:
            continue

        title = clean_title(title_match.group(1))
        overview = extract_section("Comprehensive Overview", section) or extract_section("Overview", section)
        core_concepts = extract_core_concepts(section)
        practical_applications = extract_section("Practical Applications", section)
        best_practices = extract_best_practices(section)

        topics.append({
            "title": title,
            "overview": overview,
            "coreConcepts": core_concepts,
            "practicalApplications": practical_applications,
            "bestPractices": best_practices
        })

    return {
        "chapter": {
            "topics": topics
        }
    }


# def parse_markdown_to_scorm_object(md: str):
#     import re

#     def extract_section(heading, content):
#         pattern = rf"###* {re.escape(heading)}\n+(.*?)(?=\n### |\n## |\Z)"
#         match = re.search(pattern, content, re.DOTALL)
#         return match.group(1).strip() if match else ""

#     def extract_core_concepts(section):
#         def get_text(label):
#             pattern = rf"\*\*{label}\*\*: (.*)"
#             match = re.search(pattern, section)
#             return match.group(1).strip() if match else ""

#         def get_list(label):
#             pattern = rf"\*\*{label}\*\*:\n((?:\s*-\s+.*\n?)+)"
#             match = re.search(pattern, section)
#             if not match:
#                 return []
#             items = re.findall(r"- (.+)", match.group(1))
#             return [item.strip() for item in items]

#         return {
#             "definition": get_text("Definition"),
#             "theoreticalFoundation": get_text("Theoretical Foundation"),
#             "keyComponents": get_list("Key Components")
#         }

#     def extract_best_practices(section):
#         pattern = r"\*\*Best Practices\*\*\n((?:\s*\d+\.\s+.*\n?)+)"
#         match = re.search(pattern, section)
#         return re.findall(r"\d+\.\s+(.*)", match.group(1)) if match else []

#     def extract_title(section):
#         match = re.search(r"#### (.*)", section)
#         return match.group(1).strip() if match else "Untitled"

#     topic_sections = re.split(r"\n(?=#### )", md)
#     topics = []

#     for section in topic_sections[1:]:  # skip the chapter-level intro
#         title = extract_title(section)
#         overview = extract_section("Comprehensive Overview", section)
#         core_concepts = extract_core_concepts(section)
#         practical_applications = extract_section("Practical Applications", section)
#         best_practices = extract_best_practices(section)

#         topics.append({
#             "title": title,
#             "overview": overview,
#             "coreConcepts": core_concepts,
#             "practicalApplications": practical_applications,
#             "bestPractices": best_practices
#         })

#     return {
#         "chapter": {
#             "topics": topics
#         }
#     }


# ADD these helper functions at the bottom of your routes.py file:
def calculate_materials_stats(materials):
    """Calculate statistics for materials dashboard."""
    modules = materials.get('modules', [])
    stats = {
        'total_modules': len(modules),
        'total_components': 0,
        'completion_rate': 0,
        'total_pages': 0
    }
    
    if not modules:
        return stats
    
    total_possible = len(modules) * 5  # 5 components per module
    completed = 0
    
    for module in modules:
        for component_type, component_data in module.get('components', {}).items():
            if component_data:
                stats['total_components'] += 1
                completed += 1
                
                # Estimate pages (rough calculation)
                if isinstance(component_data, dict):
                    content_str = json.dumps(component_data)
                    words = len(content_str.split())
                    stats['total_pages'] += max(1, words // 300)  # ~300 words per page
    
    stats['completion_rate'] = int((completed / total_possible) * 100) if total_possible > 0 else 0
    
    return stats

def sanitize_filename(filename):
    """Sanitize filename for safe file system usage."""
    # Remove invalid characters
    invalid_chars = '<>:"/\\|?*'
    for char in invalid_chars:
        filename = filename.replace(char, '')
    
    # Limit length
    filename = filename[:50]
    
    return filename.strip()

def create_materials_overview(analysis_data):
    """Create a text overview of all materials."""
    overview = f"""COURSE MATERIALS OVERVIEW
========================

Course: {analysis_data['course_topic']}
Audience Level: {analysis_data['audience_type'].title()}
Generated: {analysis_data.get('materials_generated_date', 'Unknown')}
Total Modules: {len(analysis_data['course_materials']['modules'])}

MODULE SUMMARY
==============

"""
    
    for module in analysis_data['course_materials']['modules']:
        overview += f"\nModule {module['number']}: {module['title']}\n"
        overview += "-" * 50 + "\n"
        
        for component_type, component_data in module['components'].items():
            if component_data:
                overview += f"? {component_type.replace('_', ' ').title()}\n"
            else:
                overview += f"? {component_type.replace('_', ' ').title()} (not generated)\n"
    
    return overview

def format_material_as_text(material_type, material_data):
    """Format material data as readable text."""
    # This is a simplified version - you can expand this based on material structure
    text = f"{material_type.replace('_', ' ').upper()}\n"
    text += "=" * 50 + "\n\n"
    
    # Remove metadata for cleaner text
    if isinstance(material_data, dict):
        data = {k: v for k, v in material_data.items() if k != 'metadata'}
    else:
        data = material_data
    
    # Convert to formatted text
    text += json.dumps(data, indent=2)
    
    return text

def format_keyword_sections(text):
    keywords = [
        'Definition',
        'Theoretical Foundation',
        'Key Components',
        'How It Works','Recommended Readings','Online Tutorials','Practice Platforms',
        'Professional Communities'
    ]

    # Match and insert \n after standalone keyword headers (not inline)
    for term in keywords:
        pattern = rf'(^[\*\+\-\s]*\**\s*{re.escape(term)}\s*\**\s*[:\-])\s*'
        text = re.sub(pattern, r'\1\n', text, flags=re.IGNORECASE | re.MULTILINE)

    # Match and insert \n after ExampleX: Something: pattern
    example_pattern = r'(^[\*\+\-\s]*\**\s*Example\s*\d+\s*:\s*[^:\n]+?:\**)\s*'
    text = re.sub(example_pattern, r'\1\n', text, flags=re.IGNORECASE | re.MULTILINE)

    # Match and insert \n after ChallengeX: Something: pattern
    challenge_pattern = r'(^[\*\+\-\s]*\**\s*Challenge\s*\d+\s*:\s*[^:\n]+?:\**)\s*'
    text = re.sub(challenge_pattern, r'\1\n', text, flags=re.IGNORECASE | re.MULTILINE)

    # Insert \n BEFORE Solution: only when it's a header or list item
    solution_pattern = r'\s*([\*\+\-\s]*\**\s*Solution\s*:\**)'
    text = re.sub(solution_pattern, r'\n\1', text, flags=re.IGNORECASE | re.MULTILINE)


    return text


import requests
from flask import request,jsonify
from bs4 import BeautifulSoup
import shutil


BASE_DIR = os.path.dirname(os.path.abspath(__file__))  # this points to /app
#SAVE_DIR = os.path.join(BASE_DIR, "static", "materialimages")
MODULE_IMAGES_BASE = os.path.join(BASE_DIR, "..", "moduleimages")  # outside static

def safe_filename(text, idx):
    """Sanitize alt text into a safe filename."""
    if not text:
        return f"image_{idx}"
    text = text.strip().replace(" ", "_")
    text = re.sub(r'[^A-Za-z0-9_\-]', '', text)   # keep only safe chars
    return text[:50]


from urllib.parse import urlparse    
def extract_and_download_figures(md_text, analysis_id):
    """
    Extract <figure> and markdown image tags from md_text,
    download images into /moduleimages/<analysis_id>/,
    and replace references with local <img> tags.
    """
    save_folder = os.path.join(MODULE_IMAGES_BASE, str(analysis_id))
    os.makedirs(save_folder, exist_ok=True)


    # --- Helper to detect if a URL is local to this backend ---
    def is_local(img_url):
        parsed = urlparse(img_url)
        # Relative URLs are local
        if not parsed.netloc:
            return True
        # Absolute URLs → check if hostname matches current backend host
        return parsed.netloc == request.host

    # --- Step 1: Convert markdown-style images to HTML ---
    md_img_pattern = re.compile(r'!\[([^\]]*)\]\(([^)]+)\)')
    
    def md_to_div(match):
        alt_text = match.group(1) or "image"
        img_url = match.group(2)

        # if is_local(img_url):
        #     # Local URL → keep as-is
        #     return f'<div class="textbook-image"><img src="{img_url}" alt="{alt_text}"/></div>'
        
        if is_local(img_url):
            # Local URL → convert to relative path inside SCORM package
            url_path = img_url.split("?")[0]
            base_name = os.path.basename(url_path)
            rel_path = os.path.join("moduleimages", str(analysis_id), base_name).replace("\\", "/")
            return f'<div class="textbook-image"><img src="{rel_path}" alt="{alt_text}"/></div>'


        # Extract the filename from the URL
        url_path = img_url.split("?")[0]  # remove query params
        base_name = os.path.basename(url_path)  # e.g., "graphical_user_interface.jpg"
        file_name = base_name or "image.jpg"
        
        # safe_name = re.sub(r'\W+', '_', alt_text.strip()).strip("_")
        # file_ext = os.path.splitext(img_url.split("?")[0])[1] or ".jpg"
        # file_name = f"{safe_name}{file_ext}"

        file_path = os.path.join(save_folder, file_name)
        rel_path = os.path.join("moduleimages", str(analysis_id), file_name).replace("\\", "/")

        if not os.path.exists(file_path):
            try:
                response = requests.get(img_url, timeout=30)
                response.raise_for_status()
                with open(file_path, "wb") as f:
                    f.write(response.content)
            except Exception as e:
                print(f"[Download Error] {img_url}: {e}")

        return f'<div class="textbook-image"><img src="{rel_path}" alt="{alt_text}"/></div>'

    md_text = md_img_pattern.sub(md_to_div, md_text)

    soup = BeautifulSoup(md_text, "html.parser")

    # --- Step 2: Handle <figure> tags and download images ---
    for idx, fig in enumerate(soup.find_all("figure"), start=1):
        try:
            img_tag = fig.find("img")
            if not img_tag:
                continue

            img_url = img_tag.get("src")
            alt_text = img_tag.get("alt", f"image_{idx}")
            

            safe_name = re.sub(r'\W+', '_', alt_text.strip()).strip("_")
            file_ext = os.path.splitext(img_url.split("?")[0])[1] or ".jpg"
            file_name = f"{safe_name}{file_ext}"
            file_path = os.path.join(save_folder, file_name)

            # Download if not exists
            if not os.path.exists(file_path):
                try:
                    response = requests.get(img_url, timeout=30)
                    response.raise_for_status()
                    with open(file_path, "wb") as f:
                        f.write(response.content)
                except Exception as e:
                    print(f"[Download Error] {img_url}: {e}")
                    continue

            # Replace <figure> with <img>
            rel_path = os.path.join("moduleimages", str(analysis_id), file_name).replace("\\", "/")
            new_img_tag = soup.new_tag("img", src=rel_path, alt=alt_text)
            fig.replace_with(new_img_tag)

        except Exception as e:
            print(f"[Figure Error] {e}")
            continue

    return str(soup)



@main.route("/upload_image", methods=["POST"])
def upload_image():
    if 'image' not in request.files:
        return jsonify({"success": False, "message": "No file provided"}), 400

    file = request.files['image']
    analysis_id = request.form.get("analysis_id", "default_analysis")

    # --- Step 1: Make folder ---
    save_folder = os.path.join(MODULE_IMAGES_BASE, str(analysis_id))
    os.makedirs(save_folder, exist_ok=True)

    # --- Step 2: Generate safe filename ---
    # alt_text = request.form.get("alt_text", os.path.splitext(file.filename)[0])
    # safe_name = re.sub(r'\W+', '_', alt_text.strip()).strip("_")
    safe_name = re.sub(r'\W+', '_', os.path.splitext(file.filename)[0])
    file_ext = os.path.splitext(file.filename)[1] or ".jpg"
    #timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    # file_name = f"{safe_name}_{timestamp}{file_ext}"
    file_name = f"{safe_name}{file_ext}"
    file_path = os.path.join(save_folder, file_name)

    # --- Step 3: Save file ---
    file.save(file_path)

    # --- Step 4: Build dynamic URL ---
    # request.host_url gives the current scheme + host + port
    rel_path = os.path.join("moduleimages", str(analysis_id), file_name).replace("\\", "/")
    file_url = f"{request.host_url.rstrip('/')}/{rel_path}"

    return jsonify({"success": True, "url": file_url})

@main.route('/process-md/<analysis_id>', methods=["POST"])
def process_md(analysis_id):
    try:
        data = request.get_json()
        # if not data or "md_text" not in data:
        #     return jsonify({"error": "md_text missing"}), 400

        # md_text = data["main_content"]

        possible_keys = [
            "main_content",
            "comprehensive_assessments",
            "comprehensive_activities",
            "comprehensive_lesson_plan",
            "comprehensive_instructor_guide",
        ]

        md_text = next((data[k] for k in possible_keys if k in data and data[k]), None)

        # Run your existing function
        updated_md_text = extract_and_download_figures(md_text,analysis_id)

        # Update the object
        data = updated_md_text

        return jsonify(data)

    except Exception as e:
        return jsonify({"error": str(e)}), 500



def parse_content_to_json_contenttype(md_text,analysis_id,coursetopic):
    # rawtext = md_text
    rawtext = format_keyword_sections(md_text)
    #imgresult = extract_and_download_figures(md_text)
    rawtext = extract_and_download_figures(rawtext,analysis_id)
    #print(imgresult)

    headings = [
        "Chapter", "Learning Outcomes", "Chapter Overview", "Introduction",
        "Detailed Topic Coverage", "Synthesis and Integration",
        "Practical Implementation Guide", "Tools and Resources", "Essential Tools",
        "Additional Resources", "Recommended Readings", "Online tutorials",
        "Practice platforms", "Professional communities", "Chapter Summary",
        "Key Terms Glossary"
    ]

    topic_headings = [
        "Comprehensive Overview", "Core Concepts", "Definition", "Theoretical Foundation",
        "Key Components", "How It Works", "Detailed Examples", "Practical Applications",
        "Common Challenges and Solutions", "Best Practices", "Integration with Other Concepts"
    ]

    def clean_line(line):
        cleaned = re.sub(r'^#+\s*', '', line)
        cleaned = re.sub(r'^[-*+]\s+', '', cleaned)
        #cleaned = re.sub(r'^\d+\.\s+', '', cleaned)
        cleaned = re.sub(r'\*\*|__|\*|_', '', cleaned)
        cleaned = re.sub(r'`+', '', cleaned)
        cleaned = re.sub(r':\s*$', '', cleaned)
        cleaned = re.sub(r'\t', ' ', cleaned)
        cleaned = re.sub(r'JAVAHOME', 'JAVA_HOME', cleaned)
        return cleaned.strip()

    def process_examples_lines(example_key, lines, current_topic):
        detailed_examples = []
        current_level = None
        current_steps = []
        for line in lines:
            line = line.strip()
            if line.startswith("Example"):
                if current_level:
                    detailed_examples.append({
                        "level": current_level,
                        "steps": [step for step in current_steps if step]
                    })
                    current_steps = []
                current_level = line.split(":", 1)[1].strip() if ":" in line else line.strip()
            elif current_level and line:
                current_steps.append(line)
        if current_level:
            detailed_examples.append({
                "level": current_level,
                "steps": [step for step in current_steps if step]
            })
        current_topic["Detailed Examples"] = detailed_examples
        return current_topic

    def process_challenge_lines(challenge_key, lines, current_topic):
        challenges_solutions = []
        current_challenge = None
        current_steps = []
        is_pattern_2 = False
        for i, line in enumerate(lines):
            line = line.strip()
            if "Challenge" in line:
                if current_challenge:
                    if is_pattern_2:
                        if current_steps and "Solution" in current_steps[0]:
                            solution = current_steps[0].split(":", 1)[1].strip() if ":" in current_steps[0] else current_steps[0].strip()
                            challenges_solutions.append({
                                "challenge": current_challenge,
                                "solution": solution
                            })
                    else:
                        challenges_solutions.append({
                            "challenge": current_challenge,
                            "solution": [step for step in current_steps if step]
                        })
                    current_steps = []
                challenge_part = line.split(":", 1)[1].strip() if ":" in line else line.strip()
                current_challenge = challenge_part.rstrip("-").strip()
                if i + 1 < len(lines) and "Solution" in lines[i + 1].strip():
                    is_pattern_2 = True
                else:
                    is_pattern_2 = False
            elif current_challenge and line:
                current_steps.append(line)
        if current_challenge:
            if is_pattern_2:
                if current_steps and "Solution" in current_steps[0]:
                    solution = current_steps[0].split(":", 1)[1].strip() if ":" in current_steps[0] else current_steps[0].strip()
                    challenges_solutions.append({
                        "challenge": current_challenge,
                        "solution": solution
                    })
            else:
                challenges_solutions.append({
                    "challenge": current_challenge,
                    "solution": [step for step in current_steps if step]
                })
        current_topic["Common Challenges and Solutions"] = challenges_solutions
        return current_topic

    def restructure_challenges(current_topic):
        challenges_solutions = current_topic["Common Challenges and Solutions"]
        restructured = []
        for challenge in challenges_solutions:
            if any("Description:" in item for item in challenge["solution"]):
                description = ""
                solution = ""
                for item in challenge["solution"]:
                    if "Description:" in item:
                        description = item.split(":", 1)[1].strip()
                    elif "Solution:" in item:
                        solution = item.split(":", 1)[1].strip()
                restructured.append({
                    "challenge": description or challenge["challenge"],
                    "solution": solution
                })
            else:
                restructured.append({
                    "challenge": challenge["challenge"],
                    "solution": challenge["solution"]
                })
        current_topic["Common Challenges and Solutions"] = restructured
        return current_topic

    import re,string

    def is_tag_line(line: str) -> bool:
        """Return True if the line starts with a tag like <div>, <img>, <figure>, <textarea>, etc."""
        stripped = line.strip().lower()
        return stripped.startswith("<div") or stripped.startswith("<img") \
            or stripped.startswith("<figure") or stripped.startswith("<textarea")

    def fix_titles_before_comprehensive(lines):
        updated_lines = lines.copy()
        comp_indices = []

        # Step 1: Identify all lines that say "Comprehensive Overview"
        for i, line in enumerate(lines):
            if line.strip().lower() == "comprehensive overview":
                comp_indices.append(i)

        # Step 2: Go to the line before each and update
        for idx, comp_idx in enumerate(comp_indices):
            prev_idx = comp_idx - 1

            # Walk upwards until a non-tag/text line is found
            
            while prev_idx >= 0 and is_tag_line(lines[prev_idx]):
                prev_idx -= 1

            if prev_idx >= 0:
                prev_line = updated_lines[prev_idx].strip()
                letter = string.ascii_uppercase[idx]
                
                # Remove any existing "X. " or "x. " prefix if present
                prev_line = re.sub(r'^[A-Za-z]\.\s*', '', prev_line).strip()
                
                # Capitalize and add the correct letter prefix
                capitalized = prev_line[0].upper() + prev_line[1:] if prev_line else "Untitled Topic"
                updated_lines[prev_idx] = f"{letter}. {capitalized}"

        return updated_lines
    def modify_detailed_topic_coverage(standard_json):
        """
        Modifies the 'Detailed Topic Coverage' section in the provided JSON by:
        1. Removing leading alphabets from 'Topic Title' (e.g., 'A.' from 'A. General Content').
        2. Removing topics with empty values for key fields.
        Updates the standard_json and returns it.
        
        Args:
            standard_json (dict): The input JSON data containing the 'chapter' and 'Detailed Topic Coverage' keys.
        
        Returns:
            dict: The updated JSON with modified 'Detailed Topic Coverage'.
        """
        try:
            # Access the Detailed Topic Coverage section
            detailed_topics = standard_json["chapter"]["Detailed Topic Coverage"]
            
            # Initialize a new list for modified topics
            modified_topics = []
            
            for topic in detailed_topics:
                # Check if the topic has non-empty values for key fields
                is_empty = (
                    not topic.get("Comprehensive Overview") and
                    not topic.get("Core Concepts", {}).get("Definition") and
                    not topic.get("Core Concepts", {}).get("Theoretical Foundation") and
                    not topic.get("Core Concepts", {}).get("Key Components") and
                    not topic.get("Core Concepts", {}).get("How It Works") and
                    not topic.get("Detailed Examples") and
                    not topic.get("Practical Applications") and
                    not topic.get("Common Challenges and Solutions") and
                    not topic.get("Best Practices") and
                    not topic.get("Integration with Other Concepts")
                )
                
                # Skip topics that are empty
                if is_empty:
                    continue
                    
                # Remove leading alphabet and period (e.g., 'A.' or 'B.') from Topic Title
                topic["Topic Title"] = re.sub(r'^[A-Z]\.\s*', '', topic["Topic Title"])
                
                # Add the modified topic to the new list
                modified_topics.append(topic)
            
            # Update the standard_json with the modified topics
            standard_json["chapter"]["Detailed Topic Coverage"] = modified_topics           
            return standard_json
        
        except KeyError as e:
            print(f"Error: Key {e} not found in JSON structure")
            return standard_json
        except Exception as e:
            print(f"An error occurred: {str(e)}")
            return standard_json

    def process_topic_coverage(topic_key, lines, standard_json):
        lines = fix_titles_before_comprehensive(lines)
        topics = []
        current_topic = {
            "Topic Title": "",
            "Comprehensive Overview": "",
            "Core Concepts": {
                "Definition": "",
                "Theoretical Foundation": "",
                "Key Components": [],
                "How It Works": []
            },
            "Detailed Examples": [],
            "Practical Applications": "",
            "Common Challenges and Solutions": [],
            "Best Practices": [],
            "Integration with Other Concepts": ""
        }
        current_subsection = None
        current_lines = []
        topic_title = None
        lowercase_topic_headings = [h.lower() for h in topic_headings]
        # for line in lines:
        for i, line in enumerate(lines):
            line = line.strip()
            if re.match(r'^[A-Z]\.\s*+.+$', line):
                if topic_title:
                    if current_subsection:
                        if current_subsection == "Comprehensive Overview":
                            joined = '\n'.join([l for l in current_lines if l])
                            current_topic["Comprehensive Overview"] = joined.split('\n') if '\n' in joined else joined
                        elif current_subsection == "Definition":
                            joined = '\n'.join([l for l in current_lines if l])
                            current_topic["Core Concepts"]["Definition"] = joined.split('\n') if '\n' in joined else joined
                        elif current_subsection == "Theoretical Foundation":
                            joined = '\n'.join([l for l in current_lines if l])
                            current_topic["Core Concepts"]["Theoretical Foundation"] = joined.split('\n') if '\n' in joined else joined
                        elif current_subsection == "Key Components":
                            current_topic["Core Concepts"]["Key Components"] = [l for l in current_lines if l]
                        elif current_subsection == "How It Works":
                            current_topic["Core Concepts"]["How It Works"] = [l for l in current_lines if l]
                        elif current_subsection == "Detailed Examples":
                            current_topic = process_examples_lines(current_subsection, current_lines, current_topic)
                        elif current_subsection == "Practical Applications":
                            joined = '\n'.join([l for l in current_lines if l])
                            current_topic["Practical Applications"] = joined.split('\n') if '\n' in joined else joined
                        elif current_subsection == "Common Challenges and Solutions":
                            current_topic = process_challenge_lines(current_subsection, current_lines, current_topic)
                            current_topic = restructure_challenges(current_topic)
                        elif current_subsection == "Best Practices":
                            current_topic["Best Practices"] = [l for l in current_lines if l]
                        elif current_subsection == "Integration with Other Concepts":
                            joined = '\n'.join([l for l in current_lines if l])
                            current_topic["Integration with Other Concepts"] = joined.split('\n') if '\n' in joined else joined
                        elif line.startswith("<div"):
                            # Add new key to store the section line
                            current_topic["Section"] = line
                    current_topic["Topic Title"] = topic_title
                    topics.append(current_topic)
                    current_topic = {
                        "Topic Title": "",
                        "Comprehensive Overview": "",
                        "Core Concepts": {
                            "Definition": "",
                            "Theoretical Foundation": "",
                            "Key Components": [],
                            "How It Works": []
                        },
                        "Detailed Examples": [],
                        "Practical Applications": "",
                        "Common Challenges and Solutions": [],
                        "Best Practices": [],
                        "Integration with Other Concepts": ""
                    }
                    current_subsection = None
                    current_lines = []
                topic_title = line
            elif line.lower() in lowercase_topic_headings:
                matching_heading = next(h for h in topic_headings if h.lower() == line.lower())
                if current_subsection:
                    if current_subsection == "Comprehensive Overview":
                        joined = '\n'.join([l for l in current_lines if l])
                        current_topic["Comprehensive Overview"] = joined.split('\n') if '\n' in joined else joined
                    elif current_subsection == "Definition":
                        joined = '\n'.join([l for l in current_lines if l])
                        current_topic["Core Concepts"]["Definition"] = joined.split('\n') if '\n' in joined else joined
                    elif current_subsection == "Theoretical Foundation":
                        joined = '\n'.join([l for l in current_lines if l])
                        current_topic["Core Concepts"]["Theoretical Foundation"] = joined.split('\n') if '\n' in joined else joined
                    elif current_subsection == "Key Components":
                        current_topic["Core Concepts"]["Key Components"] = [l for l in current_lines if l]
                    elif current_subsection == "How It Works":
                        current_topic["Core Concepts"]["How It Works"] = [l for l in current_lines if l]
                    elif current_subsection == "Detailed Examples":
                        current_topic = process_examples_lines(current_subsection, current_lines, current_topic)
                    elif current_subsection == "Practical Applications":
                        joined = '\n'.join([l for l in current_lines if l])
                        current_topic["Practical Applications"] = joined.split('\n') if '\n' in joined else joined
                    elif current_subsection == "Common Challenges and Solutions":
                        current_topic = process_challenge_lines(current_subsection, current_lines, current_topic)
                        current_topic = restructure_challenges(current_topic)
                    elif current_subsection == "Best Practices":
                        current_topic["Best Practices"] = [l for l in current_lines if l]
                    elif current_subsection == "Integration with Other Concepts":
                        joined = '\n'.join([l for l in current_lines if l])
                        current_topic["Integration with Other Concepts"] = joined.split('\n') if '\n' in joined else joined
                    elif line.startswith("<div"):
                        # Add new key to store the section line
                        current_topic["Section"] = line
                current_subsection = matching_heading
                current_lines = []
            elif line.startswith("<div") and i > 0 and re.match(r"^[A-Za-z]\.", lines[i - 1].strip()):
                current_topic["Section"] = line
            elif current_subsection and line:
                current_lines.append(line)
        if topic_title:
            if current_subsection:
                if current_subsection == "Comprehensive Overview":
                    joined = '\n'.join([l for l in current_lines if l])
                    current_topic["Comprehensive Overview"] = joined.split('\n') if '\n' in joined else joined
                elif current_subsection == "Definition":
                    joined = '\n'.join([l for l in current_lines if l])
                    current_topic["Core Concepts"]["Definition"] = joined.split('\n') if '\n' in joined else joined
                elif current_subsection == "Theoretical Foundation":
                    joined = '\n'.join([l for l in current_lines if l])
                    current_topic["Core Concepts"]["Theoretical Foundation"] = joined.split('\n') if '\n' in joined else joined
                elif current_subsection == "Key Components":
                    current_topic["Core Concepts"]["Key Components"] = [l for l in current_lines if l]
                elif current_subsection == "How It Works":
                    current_topic["Core Concepts"]["How It Works"] = [l for l in current_lines if l]
                elif current_subsection == "Detailed Examples":
                    current_topic = process_examples_lines(current_subsection, current_lines, current_topic)
                elif current_subsection == "Practical Applications":
                    joined = '\n'.join([l for l in current_lines if l])
                    current_topic["Practical Applications"] = joined.split('\n') if '\n' in joined else joined
                elif current_subsection == "Common Challenges and Solutions":
                    current_topic = process_challenge_lines(current_subsection, current_lines, current_topic)
                    current_topic = restructure_challenges(current_topic)
                elif current_subsection == "Best Practices":
                    current_topic["Best Practices"] = [l for l in current_lines if l]
                elif current_subsection == "Integration with Other Concepts":
                    joined = '\n'.join([l for l in current_lines if l])
                    current_topic["Integration with Other Concepts"] = joined.split('\n') if '\n' in joined else joined
                elif line.startswith("<div"):
                    # Add new key to store the section line
                    current_topic["Section"] = line
            current_topic["Topic Title"] = topic_title
            topics.append(current_topic)
        standard_json["chapter"]["Detailed Topic Coverage"] = topics
        check_missing_values(standard_json)
        if topics == []:
            standard_json = process_topic_coverage_two(topic_key, lines, standard_json)
        updated_json = modify_detailed_topic_coverage(standard_json)
        return updated_json                                                                                      
        #return standard_json
        
    
    
    def check_missing_values(standard_json, parent_key="", chapter_key="chapter"):
        """
        Recursively checks the standard_json for empty or missing values and logs them.
        
        Args:
            standard_json (dict): The JSON structure to check.
            parent_key (str): The parent key path for nested properties (used for logging).
            chapter_key (str): The key for the chapter in the JSON (default: 'chapter').
        
        Returns:
            None: Logs missing/empty values using current_app.logger.info.
        """
        def is_empty(value):
            """Helper function to check if a value is considered empty."""
            if value is None:
                return True
            if isinstance(value, str) and value.strip() == "":
                return True
            if isinstance(value, (list, dict)) and len(value) == 0:
                return True
            return False

        def recursive_check(obj, current_path):
            """Recursively checks nested JSON structure for empty values."""
            if isinstance(obj, dict):
                for key, value in obj.items():
                    new_path = f"{current_path}.{key}" if current_path else key
                    if isinstance(value, (dict, list)):
                        if is_empty(value):
                            current_app.logger.info(f"Missing or empty value for property: {new_path}")
                        else:
                            recursive_check(value, new_path)
                    elif is_empty(value):
                        current_app.logger.info(f"Missing or empty value for property: {new_path}")
            elif isinstance(obj, list):
                for i, item in enumerate(obj):
                    new_path = f"{current_path}[{i}]"
                    if isinstance(item, (dict, list)):
                        if is_empty(item):
                            current_app.logger.info(f"Missing or empty value for property: {new_path}")
                        else:
                            recursive_check(item, new_path)
                    elif is_empty(item):
                        current_app.logger.info(f"Missing or empty value for property: {new_path}")

        # Start checking from the chapter level
        if chapter_key in standard_json:
            recursive_check(standard_json[chapter_key], chapter_key)
        else:
            current_app.logger.info(f"Missing chapter key: {chapter_key}")

    def process_subsection(subsection, lines, current_topic):
        # Skip if no content
        if not lines:
            return
       
        content = '\n'.join(lines)
       
        if subsection == "Comprehensive Overview":
            current_topic["Comprehensive Overview"] = content
       
        elif subsection == "Definition":
            current_topic["Core Concepts"]["Definition"] = content
       
        elif subsection == "Theoretical Foundation":
            current_topic["Core Concepts"]["Theoretical Foundation"] = content
       
        elif subsection == "Key Components":
            current_topic["Core Concepts"]["Key Components"] = [line.strip() for line in lines if line.strip()]
       
        elif subsection == "How It Works":
            current_topic["Core Concepts"]["How It Works"] = [line.strip() for line in lines if line.strip()]
       
        elif subsection == "Detailed Examples":
            process_examples_lines(subsection, lines, current_topic)
       
        elif subsection == "Practical Applications":
            current_topic["Practical Applications"] = content
       
        elif subsection == "Common Challenges and Solutions":
            process_challenge_lines(subsection, lines, current_topic)
            restructure_challenges(current_topic)
       
        elif subsection == "Best Practices":
            current_topic["Best Practices"] = [line.strip() for line in lines if line.strip()]
       
        elif subsection == "Integration with Other Concepts":
            current_topic["Integration with Other Concepts"] = content
 
    def process_topic_coverage_two(topic_key, lines, standard_json):
        topics = []
        current_topic = {
            "Topic Title": "",
            "Comprehensive Overview": "",
            "Core Concepts": {
                "Definition": "",
                "Theoretical Foundation": "",
                "Key Components": [],
                "How It Works": []
            },
            "Detailed Examples": [],
            "Practical Applications": "",
            "Common Challenges and Solutions": [],
            "Best Practices": [],
            "Integration with Other Concepts": ""
        }
        current_subsection = None
        current_lines = []
        lowercase_topic_headings = [h.lower() for h in topic_headings]
 
        for line in lines:
            line_lower = line.strip().lower()
            if line_lower in lowercase_topic_headings:
                # Process accumulated lines for the previous subsection
                if current_subsection:
                    process_subsection(current_subsection, current_lines, current_topic)
                    current_lines = []
                current_subsection = next(h for h in topic_headings if h.lower() == line_lower)
            else:
                if not current_topic["Topic Title"] and line.strip():
                    # First valid line becomes the topic title
                    current_topic["Topic Title"] = line
                elif current_topic["Topic Title"] and not current_subsection and line.strip():
                    # New topic detected: finalize current and start new
                    topics.append(current_topic)
                    current_topic = { ... }  # Reset structure (same as above)
                    current_topic["Topic Title"] = line
                elif current_subsection and line.strip():
                    # Accumulate content under the current subsection
                    current_lines.append(line)
 
        # Process the last subsection and topic
        if current_subsection:
            process_subsection(current_subsection, current_lines, current_topic)
        if current_topic["Topic Title"]:
            topics.append(current_topic)
       
        standard_json["chapter"]["Detailed Topic Coverage"] = topics
        return standard_json

    def process_glossary_lines(lines):
        glossary = []
        for line in lines:
            line = clean_line(line)
            if ":" in line:
                term, definition = line.split(":", 1)
                glossary.append({
                    "term": term.strip(),
                    "definition": definition.strip()
                })
        return glossary

    raw_lines = [line.strip() for line in rawtext.split('\n') if line.strip()]
    #cleaned_lines = [clean_line(line) for line in raw_lines]
    cleaned_lines = [
    line.strip() if "<img" in line.lower() else clean_line(line)
    for line in raw_lines
    ]



    parsed_data = {}
    current_heading = None
    lowercase_headings = [h.lower() for h in headings]

    for line in cleaned_lines:
        if re.search(r"<\s*img\b", line, re.IGNORECASE):
            cleaned = line.strip()   # keep <img> line unchanged
        else:
            cleaned = clean_line(line)
        if cleaned.lower() in lowercase_headings:
            matching_heading = next(h for h in headings if h.lower() == cleaned.lower())
            current_heading = matching_heading
            if current_heading not in parsed_data:
                parsed_data[current_heading] = []
        elif current_heading:
            parsed_data[current_heading].append(cleaned)

    standard_json = {
        "chapter": {
            "Chapter": "{}",
            "Learning Outcomes": "",
            "Chapter Overview": "",
            "Introduction": "",
            "Detailed Topic Coverage": [],
            "Synthesis and Integration": "",
            "Practical Implementation Guide": "",
            "Tools and Resources": {
                "Essential Tools": "",
                "Additional Resources": {
                    "Recommended Readings": "",
                    "Online tutorials": "",
                    "Practice platforms": "",
                    "Professional communities": ""
                }
            },
            "Chapter Summary": "",
            "Key Terms Glossary": []
        }
    }

    if cleaned_lines:
        standard_json["chapter"]['Chapter'] = cleaned_lines[0].split(':', 1)[1].strip() if ':' in cleaned_lines[0] else cleaned_lines[0]

    for key, value in parsed_data.items():
        joined_value = '\n'.join(value)
        final_value = joined_value.split('\n') if '\n' in joined_value else joined_value
        matching_key = next(h for h in headings if h.lower() == key.lower())
        if matching_key == "Learning Outcomes":
            standard_json["chapter"][matching_key] = final_value
        elif matching_key == "Practical Implementation Guide":
            standard_json["chapter"][matching_key] = final_value
        elif matching_key == "Chapter":
            standard_json["chapter"][matching_key] = final_value
        elif matching_key == "Essential Tools":
            standard_json["chapter"]["Tools and Resources"][matching_key] = final_value
        elif matching_key in ["Recommended Readings", "Online tutorials", "Practice platforms", "Professional communities"]:
            standard_json["chapter"]["Tools and Resources"]["Additional Resources"][matching_key] = final_value
        elif matching_key == "Detailed Topic Coverage":
            standard_json = process_topic_coverage(matching_key, value, standard_json)
        elif matching_key == "Key Terms Glossary":
            standard_json["chapter"][matching_key] = process_glossary_lines(value)
        elif matching_key in standard_json["chapter"] and not isinstance(standard_json["chapter"][matching_key], (dict, list)):
            standard_json["chapter"][matching_key] = final_value

    
    # return standard_json
    updated_json = update_standard_json(standard_json)
    return updated_json


def update_standard_json(standard_json):
    key_mapping = {
        "Chapter": "title",
        "Learning Outcomes": "learningOutcomes",
        "Chapter Overview": "overview",
        "Introduction": "introduction",
        "Detailed Topic Coverage": "topics",
        "Topic Title": "title",
        "Comprehensive Overview": "overview",
        "Core Concepts": "coreConcepts",
        "Definition": "definition",
        "Theoretical Foundation": "theoreticalFoundation",
        "Key Components": "keyComponents",
        "How It Works": "howitWorks",
        "Detailed Examples": "examples",
        "Practical Applications": "practicalApplications",
        "Common Challenges and Solutions": "challengesAndSolutions",
        "Best Practices": "bestPractices",
        "Synthesis and Integration": "synthesis",
        "Practical Implementation Guide": "implementationGuide",
        "Tools and Resources": "toolsAndResources",
        "Essential Tools": "essentialTools",
        "Additional Resources": "additionalResources",
        "Recommended Readings": "recommendedReadings",
        "Online tutorials": "onlineTutorials",
        "Practice platforms": "practicePlatforms",
        "Professional communities": "professionalCommunities",
        "Chapter Summary": "summary",
        "Key Terms Glossary": "glossary"
    }

    return rename_keys(standard_json, key_mapping)

def rename_keys(obj, key_mapping):
    if isinstance(obj, dict):
        new_obj = {}
        for k, v in obj.items():
            new_key = key_mapping.get(k, k)
            new_obj[new_key] = rename_keys(v, key_mapping)
        return new_obj
    elif isinstance(obj, list):
        return [rename_keys(item, key_mapping) for item in obj]
    else:
        return obj

@main.route('/download_module_materials/<analysis_id>/<int:module_id>')
def download_module_materials(analysis_id, module_id):
    """Download materials for a specific module."""
    current_app.logger.info("Preparing download for module %s (analysis ID: %s)", module_id, analysis_id)

    # Load analysis data
    analysis_data = load_analysis(analysis_id)
    
    if not analysis_data or 'course_materials' not in analysis_data:
        current_app.logger.warning("Materials not found for analysis ID: %s", analysis_id)
        flash('Materials not found.')
        return redirect(url_for('main.index'))
    
    # Find the module
    module = None
    for m in analysis_data['course_materials']['modules']:
        if m['number'] == module_id:
            module = m
            break
    
    if not module:
        current_app.logger.warning("Module %s not found in analysis ID: %s", module_id, analysis_id)
        flash('Module not found.')
        return redirect(url_for('main.view_materials', analysis_id=analysis_id))
    
    try:
        current_app.logger.debug("Creating ZIP for module %s", module_id)
        # Create a zip file with module materials
        zip_buffer = io.BytesIO()
        
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            # Add each component as a JSON file
            for component_type, component_data in module.get('components', {}).items():
                if component_data:
                    filename = f"Module_{module_id}_{component_type}.json"
                    zip_file.writestr(filename, json.dumps(component_data, indent=2))
            
            # Add a summary file
            summary = {
                'module_number': module['number'],
                'module_title': module['title'],
                'generated_components': list(module.get('components', {}).keys()),
                'course_topic': analysis_data['course_topic'],
                'generation_date': analysis_data.get('materials_generated_date', 'Unknown')
            }
            zip_file.writestr(f"Module_{module_id}_Summary.json", json.dumps(summary, indent=2))
        
        zip_buffer.seek(0)
        current_app.logger.info("Module ZIP created successfully for module %s", module_id)
        
        return send_file(
            zip_buffer,
            mimetype='application/zip',
            as_attachment=True,
            download_name=f"Module_{module_id}_Materials.zip"
        )
        
    except Exception as e:
        current_app.logger.error("Error creating download file (download_module_materials): %s", str(e), exc_info=True)
        logger.error(f"Error creating module materials ZIP: {str(e)}")
        flash('Error creating download file.')
        return redirect(url_for('main.view_materials', analysis_id=analysis_id))

@main.route('/export_materials/<analysis_id>/<format>')
def export_materials(analysis_id, format):
    """Export materials in different formats (PDF, DOCX)."""
    # For now, redirect to download all materials
    # You can implement specific format exports later
    flash(f'{format.upper()} export not yet implemented. Downloading as ZIP instead.')
    return redirect(url_for('main.download_all_materials', analysis_id=analysis_id))




from flask import current_app, send_from_directory,  abort

@main.route("/moduleimages/<analysis_id>/<filename>")
def serve_module_image(analysis_id, filename):
    folder = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "moduleimages", analysis_id)
    folder = os.path.abspath(folder)  # ensure safe absolute path
    if not os.path.exists(os.path.join(folder, filename)):
        abort(404)
    return send_from_directory(folder, filename)



def calculate_materials_stats(materials):
    """Calculate statistics for materials dashboard."""
    modules = materials.get('modules', [])
    stats = {
        'total_modules': len(modules),
        'total_components': 0,
        'completion_rate': 0,
        'total_pages': 0
    }
    
    if not modules:
        return stats
    
    total_possible = len(modules) * 5  # 5 components per module
    completed = 0
    
    for module in modules:
        for component_type, component_data in module.get('components', {}).items():
            if component_data:
                stats['total_components'] += 1
                completed += 1
                
                # Estimate pages (rough calculation)
                if isinstance(component_data, dict):
                    content_str = json.dumps(component_data)
                    words = len(content_str.split())
                    stats['total_pages'] += max(1, words // 300)  # ~300 words per page
    
    stats['completion_rate'] = int((completed / total_possible) * 100) if total_possible > 0 else 0
    
    return stats

def sanitize_filename(filename):
    """Sanitize filename for safe file system usage."""
    # Remove invalid characters
    invalid_chars = '<>:"/\\|?*'
    for char in invalid_chars:
        filename = filename.replace(char, '')
    
    # Limit length
    filename = filename[:50]
    
    return filename.strip()

def create_materials_overview(analysis_data):
    """Create a text overview of all materials."""
    overview = f"""COURSE MATERIALS OVERVIEW
========================

Course: {analysis_data['course_topic']}
Audience Level: {analysis_data['audience_type'].title()}
Generated: {analysis_data.get('materials_generated_date', 'Unknown')}
Total Modules: {len(analysis_data['course_materials']['modules'])}

MODULE SUMMARY
==============

"""
    
    for module in analysis_data['course_materials']['modules']:
        overview += f"\nModule {module['number']}: {module['title']}\n"
        overview += "-" * 50 + "\n"
        
        for component_type, component_data in module['components'].items():
            if component_data:
                overview += f"? {component_type.replace('_', ' ').title()}\n"
            else:
                overview += f"? {component_type.replace('_', ' ').title()} (not generated)\n"
    
    return overview

def format_material_as_text(material_type, material_data):
    """Format material data as readable text."""
    # This is a simplified version - you can expand this based on material structure
    text = f"{material_type.replace('_', ' ').upper()}\n"
    text += "=" * 50 + "\n\n"
    
    # Remove metadata for cleaner text
    if isinstance(material_data, dict):
        data = {k: v for k, v in material_data.items() if k != 'metadata'}
    else:
        data = material_data
    
    # Convert to formatted text
    text += json.dumps(data, indent=2)
    
    return text

def create_scorm_manifest(course_title, modules):
    print(course_title)
    import xml.etree.ElementTree as ET
    from xml.dom import minidom

    def sanitize_filename(name):
        return ''.join(c if c.isalnum() else '_' for c in name)

    manifest = ET.Element('manifest', {
        'identifier': 'AIDA_COURSE_MANIFEST',
        'version': '1.0',
        'xmlns': 'http://www.imsglobal.org/xsd/imscp_v1p1',
        'xmlns:adlcp': 'http://www.adlnet.org/xsd/adlcp_v1p3',
        'xmlns:adlseq': 'http://www.adlnet.org/xsd/adlseq_v1p3',
        'xmlns:adlnav': 'http://www.adlnet.org/xsd/adlnav_v1p3',
        'xmlns:xsi': 'http://www.w3.org/2001/XMLSchema-instance',
        'xsi:schemaLocation': 'http://www.imsglobal.org/xsd/imscp_v1p1 imscp_v1p1.xsd '
                              'http://www.adlnet.org/xsd/adlcp_v1p3 adlcp_v1p3.xsd '
                              'http://www.adlnet.org/xsd/adlseq_v1p3 adlseq_v1p3.xsd '
                              'http://www.adlnet.org/xsd/adlnav_v1p3 adlnav_v1p3.xsd'
    })

    metadata = ET.SubElement(manifest, 'metadata')
    ET.SubElement(metadata, 'schema').text = 'ADL SCORM'
    ET.SubElement(metadata, 'schemaversion').text = '2004 4th Edition'

    organizations = ET.SubElement(manifest, 'organizations', {'default': 'AIDA_ORG'})
    organization = ET.SubElement(organizations, 'organization', {'identifier': 'AIDA_ORG'})
    ET.SubElement(organization, 'title').text = course_title

    root_item = ET.SubElement(organization, 'item', {
        'identifier': 'ROOT',
        'identifierref': 'RES_MAIN',
        'isvisible': 'true'
    })
    ET.SubElement(root_item, 'title').text = course_title

    # Add each module
    for mod_idx, module in enumerate(modules):
        sanitized_title = sanitize_filename(module['title'])
        folder_name = f"Module_{module['number']}_{sanitized_title}"

        mod_item = ET.SubElement(root_item, 'item', {
            'identifier': f'MODULE_{mod_idx}',
            'identifierref': f'RES_MODULE_{mod_idx}',
            'isvisible': 'true'
        })
        ET.SubElement(mod_item, 'title').text = module['title']

    # Resources section
    resources = ET.SubElement(manifest, 'resources')

    # Main resource entry
    main_resource = ET.SubElement(resources, 'resource', {
        'identifier': 'RES_MAIN',
        'type': 'webcontent',
        'adlcp:scormType': 'sco',
        'href': 'index_lms.html'
    })
    for file in ['index_lms.html', 'ADLwrapper.js', 'course.js', 'navigation.json']:
        ET.SubElement(main_resource, 'file', {'href': file})

    # Add each module resource
    for mod_idx, module in enumerate(modules):
        sanitized_title = sanitize_filename(module['title'])
        folder_name = f"Module_{module['number']}_{module['title']}"

        mod_resource = ET.SubElement(resources, 'resource', {
            'identifier': f'RES_MODULE_{mod_idx}',
            'type': 'webcontent',
            'adlcp:scormType': 'sco',
            'href': f'index_lms.html?module={mod_idx}'
        })
        ET.SubElement(mod_resource, 'file', {'href': f"{folder_name}/Content.json.clean.json"})

    # Return pretty XML string
    rough_xml = ET.tostring(manifest, 'utf-8')
    parsed_xml = minidom.parseString(rough_xml)
    return parsed_xml.toprettyxml(indent="  ")
